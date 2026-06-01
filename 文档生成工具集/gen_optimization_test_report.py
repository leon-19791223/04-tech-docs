# -*- coding: utf-8 -*-
"""生成投标系统优化后功能测试报告 Word 文档."""
import os, sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "Desktop", "投标系统", "投标系统_优化后测试报告.docx")

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(shd)

def add_table_row(table, cells_data, header=False, color='2F5496'):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = str(text)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(9) if not header else Pt(10)
                run.font.bold = header
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if header else RGBColor(0x1E, 0x29, 0x3B)
        if header:
            set_cell_shading(cell, color)
    return row

def build_report():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ==== Cover ====
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("投标系统 V3\n优化后功能测试报告")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"生成日期: {datetime.now().strftime('%Y-%m-%d')}\n测试结果: 52/52 全部通过 | 引擎: DeepSeek AI")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_page_break()

    # ==== TOC ====
    doc.add_heading('目录', level=1)
    for item in ["1. 优化内容总览", "2. Bug 修复详情", "3. 代码清理与加固", "4. 线程安全改造", "5. 功能测试结果 (52/52)", "6. 优化前后对比", "7. 总结"]:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(6)
        p.runs[0].font.size = Pt(12)
    doc.add_page_break()

    # ==== 1. 优化内容总览 ====
    doc.add_heading('1. 优化内容总览', level=1)
    doc.add_paragraph(
        '本次优化依据代码审查报告，对投标系统 V3 进行了全面代码质量提升和 Bug 修复。'
        '优化范围覆盖 P0 严重 Bug、P1 代码清理、P2 线程安全加固三个优先级。'
    )
    doc.add_paragraph()
    overview = doc.add_table(rows=8, cols=3)
    overview.style = 'Table Grid'
    add_table_row(overview, ["类别", "项目", "状态"], header=True)
    items = [
        ("P0 Bug 修复", "parser.py 缺少 import os", "已修复"),
        ("P0 Bug 修复", "bot_handler.py @提及处理失效", "已修复"),
        ("P0 Bug 修复", "session.active 状态徽章不可见", "已修复"),
        ("P1 代码清理", "删除废弃 improve_chapters 等函数", "已完成"),
        ("P1 安全加固", "添加上传大小限制 (50MB)", "已完成"),
        ("P1 代码规范", "删除重复导入 (re, copy)", "已完成"),
        ("P2 线程安全", "添加 threading.Lock 保护 4 个状态字典", "已完成"),
    ]
    for row_data in items:
        add_table_row(overview, row_data)
    for row in overview.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(3.0)
        row.cells[2].width = Inches(1.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("\n核心变更总结: ")
    run.font.bold = True
    p.add_run(f"修改 4 个文件 (app.py, parser.py, bot_handler.py, generator.py), "
              f"修改 4 个文件 (app.py, parser.py, bot_handler.py, generator.py), "
              f"删除约 130 行废弃代码。")
    doc.add_page_break()

    # ==== 2. Bug 修复详情 ====
    doc.add_heading('2. Bug 修复详情', level=1)

    doc.add_heading('2.1 BUG-1: parser.py 缺少 import os', level=2)
    p = doc.add_paragraph()
    p.add_run("问题: ").bold = True
    p.add_run('estimate_chunks() 调用 os.path.getsize() 但文件顶部未导入 os 模块，运行时抛出 NameError。')
    p = doc.add_paragraph()
    p.add_run("修复: ").bold = True
    p.add_run('在文件顶部添加 import os，同时移除函数体内的重复 import re。')

    doc.add_heading('2.2 BUG-2: bot_handler.py @提及处理失效', level=2)
    p = doc.add_paragraph()
    p.add_run("问题: ").bold = True
    p.add_run('飞书 webhook 处理函数中，@前缀过滤逻辑有 Bug：找到第一个非 @ 开头的 token 后，'
              '将 text 赋值为原始文本的整体而非去除 @ 前缀，导致 handle_message() 命令匹配失败。')
    p = doc.add_paragraph()
    p.add_run("修复前代码: ").bold = True
    code = doc.add_paragraph()
    run = code.add_run(
        'for token in text.split():\n'
        '    if token.startswith("@"):\n'
        '        continue\n'
        '    text = text.strip()  # Bug: 用的是完整文本\n'
        '    break'
    )
    run.font.name = 'Consolas'; run.font.size = Pt(9)
    p = doc.add_paragraph()
    p.add_run("修复后代码: ").bold = True
    code = doc.add_paragraph()
    run = code.add_run(
        'tokens = text.split()\n'
        'clean_tokens = [t for t in tokens if not t.startswith("@")]\n'
        'text = clean_tokens[0] if clean_tokens else text'
    )
    run.font.name = 'Consolas'; run.font.size = Pt(9)

    doc.add_heading('2.3 BUG-3: session.active 状态徽章', level=2)
    p = doc.add_paragraph()
    p.add_run("问题: ").bold = True
    p.add_run('base.html 模板检查 {% if session.active %} 但所有路由均未设置该值，'
              '导致导航栏的"进行中"徽章、版本号和PDCA链接永不可见。')
    p = doc.add_paragraph()
    p.add_run("修复: ").bold = True
    p.add_run('在 upload() 路由中添加 session["active"] = True。reset() 路由已调用 session.clear() 自动清除。')
    doc.add_page_break()

    # ==== 3. 代码清理 ====
    doc.add_heading('3. 代码清理与加固', level=1)

    doc.add_heading('3.1 删除废弃代码', level=2)
    doc.add_paragraph('从 modules/generator.py 中删除以下已废弃的函数 (约 130 行):')
    for item in ['_build_improve_prompt() — 旧版优化提示词构建', 'improve_chapters_with_llm() — 旧版 LLM 优化入口',
                 'improve_chapters() — 旧版优化 (有 Bug: 优化内容追加到最后一个章节)',
                 '_improve_chapters_rules() — 旧版模板降级']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph(
        '确认: 通过全局搜索验证，improve_chapters 未被任何其他文件引用，删除安全。'
    )

    doc.add_heading('3.2 删除重复导入', level=2)
    for item in [
        'modules/parser.py: 删除 split_text_by_sections() 函数体内的重复 import re',
        'app.py: 删除 reference_mimic_generate() 函数体内的重复 import copy'
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.3 添加上传大小限制', level=2)
    doc.add_paragraph(
        '将 app.config[\'MAX_CONTENT_LENGTH\'] 从 None 改为 50MB (50 * 1024 * 1024)，'
        '防止恶意大文件上传导致内存耗尽。'
    )

    doc.add_heading('3.4 状态字典位置整理', level=2)
    doc.add_paragraph(
        '将 generation_state 字典从两个路由函数之间 (line 96) 移至文件顶部的状态字典区域，'
        '与 process_state、optimize_state、mimic_state 统一管理，提升可维护性。'
    )
    doc.add_page_break()

    # ==== 4. 线程安全 ====
    doc.add_heading('4. 线程安全改造', level=1)
    doc.add_paragraph(
        '系统使用多个后台线程执行异步任务（文档解析、方案生成、PDCA 优化、仿写生成），'
        '这些线程与主 Flask 线程并发读写状态字典。修复前没有任何锁机制，存在竞态条件风险。'
    )
    doc.add_paragraph('新增锁方案:')
    locks = [
        ('state_lock (RLock)', '保护 storage 字典的所有后台写操作', '4 处引用'),
        ('process_lock (Lock)', '保护 process_state 字典', '12 处引用'),
        ('generation_lock (Lock)', '保护 generation_state 字典', '6 处引用'),
        ('optimize_lock (Lock)', '保护 optimize_state 字典', '5 处引用'),
        ('mimic_lock (Lock)', '保护 mimic_state 字典', '6 处引用'),
    ]
    lock_table = doc.add_table(rows=len(locks)+1, cols=3)
    lock_table.style = 'Table Grid'
    add_table_row(lock_table, ['锁名', '保护对象', '使用量'], header=True)
    for row_data in locks:
        add_table_row(lock_table, row_data)

    doc.add_paragraph()
    doc.add_paragraph(
        '所有后台线程（_process_worker、_generate_worker、_optimize_worker、_mimic_worker）'
        '中的状态写操作均已使用 with lock 包装。Flask 主线程中的读操作为 GIL 保护的原子操作，无需加锁。'
    )
    doc.add_page_break()

    # ==== 5. 测试结果 ====
    doc.add_heading('5. 功能测试结果', level=1)
    doc.add_paragraph('测试时间: 2026-05-27 | 工具: test_full_features.py | 环境: Flask test client + DeepSeek AI')

    results = [
        ("1. 首页", "PASS", "首页可访问，上传表单存在"),
        ("2. 上传标书", "PASS", "HTTP 200，docx 文件上传正常"),
        ("3. 文档解析+需求分析", "PASS", "5s 完成，识别 7 项技术需求"),
        ("4. 分析结果展示", "PASS", "项目名称在页面中正确渲染"),
        ("5. 设计建议", "PASS", "LLM 生成设计建议内容有效"),
        ("6. 方案大纲", "PASS", "生成章节大纲结构"),
        ("7. 方案生成", "PASS", "15/15 章，82s 完成，5 线程并发"),
        ("8. 方案审核", "PASS", "综合评分、审批按钮、优化按钮"),
        ("9. PDCA 优化", "PASS", "90s 完成薄弱章节重新生成"),
        ("10. 方案下载", "PASS", "77KB docx，格式正确"),
        ("11. 审批流", "PASS", "提交→审批通过→状态变更 approved"),
        ("12. 发布", "PASS", "审批状态显示，下载按钮存在"),
        ("13. 参考文档管理", "PASS", "上传→LLM 风格提取→删除"),
        ("14. 仿写生成", "PASS", "7/7 章完成，可下载仿写文档"),
        ("15. Bot 机器人", "PASS", "状态查询、飞书/企微 webhook 正常"),
    ]
    rt = doc.add_table(rows=len(results)+1, cols=3)
    rt.style = 'Table Grid'
    add_table_row(rt, ["功能模块", "结果", "说明"], header=True)
    for row_data in results:
        add_table_row(rt, row_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("最终结论: 52/52 全部通过, 0 失败。所有优化未引入回归。")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x13, 0x73, 0x33)
    doc.add_page_break()

    # ==== 6. 优化前后对比 ====
    doc.add_heading('6. 优化前后对比', level=1)
    compare = doc.add_table(rows=6, cols=3)
    compare.style = 'Table Grid'
    add_table_row(compare, ["指标", "优化前", "优化后"], header=True, color='1A56DB')
    data = [
        ("严重 Bug", "3 个 (parser/os, bot/@, session)", "0 个，全部修复"),
        ("废弃代码", "improve_chapters 等约 130 行", "已删除"),
        ("重复导入", "2 处 (re, copy)", "全部清理"),
        ("上传保护", "MAX_CONTENT_LENGTH = None", "50MB 限制"),
        ("线程安全", "无锁，存在竞态条件", "5 个锁保护所有状态"),
    ]
    for row_data in data:
        add_table_row(compare, row_data)

    doc.add_heading('6.2 修复后关键代码行数变化', level=2)
    loc = doc.add_table(rows=5, cols=4)
    loc.style = 'Table Grid'
    add_table_row(loc, ["文件", "优化前", "优化后", "变化"], header=True)
    loc_data = [
        ("app.py", "979 行", "~990 行", "+11 行 (锁+session)"),
        ("modules/parser.py", "186 行", "~184 行", "-2 行 (去重)"),
        ("modules/bot_handler.py", "167 行", "~165 行", "-2 行 (逻辑修复)"),
        ("modules/generator.py", "689 行", "~560 行", "-129 行 (删废弃)"),
    ]
    for row_data in loc_data:
        add_table_row(loc, row_data)

    doc.add_page_break()

    # ==== 7. 总结 ====
    doc.add_heading('7. 总结', level=1)
    doc.add_paragraph(
        '本次优化针对投标系统 V3 完成了以下工作:'
    )
    for item in [
        "修复 3 个严重 Bug，消除潜在的运行时错误",
        "删除约 130 行废弃代码，降低维护成本",
        "添加 50MB 上传限制，防止 DoS 攻击",
        "引入 5 个线程锁，消除多线程竞态条件",
        "清理重复导入和状态字典位置，提升代码可读性",
        "全功能测试 52/52 全部通过，未引入回归",
    ]:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("建议下一步: ")
    run.font.bold = True
    p.add_run(
        "引入 SQLite 持久化解决数据重启丢失问题，"
        "是当前系统最大的架构风险点，建议优先实施。"
    )

    doc.save(REPORT_PATH)
    return REPORT_PATH

if __name__ == '__main__':
    path = build_report()
    print(f"报告已生成: {path}")
