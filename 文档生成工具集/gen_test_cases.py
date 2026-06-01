# -*- coding: utf-8 -*-
"""生成投标系统测试用例 Word 文档（供后续测试人员使用）"""
import os, sys
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装 python-docx"); sys.exit(1)

RPATH = os.path.join(os.path.dirname(__file__), "..", "..", "Desktop", "投标系统", "投标系统_测试用例.docx")

def shd(cell, clr):
    tc = cell._element.get_or_add_tcPr()
    sh = tc.makeelement(qn('w:shd'), {qn('w:fill'): clr, qn('w:val'): 'clear'})
    tc.append(sh)

def arow(t, d, h=False, clr='2F5496'):
    r = t.add_row()
    for i, v in enumerate(d):
        c = r.cells[i]; c.text = str(v)
        for p in c.paragraphs:
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            for rn in p.runs:
                rn.font.size = Pt(8) if not h else Pt(9); rn.font.bold = h
                if h: rn.font.color.rgb = RGBColor(255,255,255)
        if h: shd(c, clr)

def build():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = '微软雅黑'; s.font.size = Pt(10)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for _ in range(4): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = 1
    r = t.add_run("投标系统 V3\n功能测试用例"); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    doc.add_paragraph()
    sub = doc.add_paragraph(); sub.alignment = 1
    r = sub.add_run(f"版本: V3.0 | 日期: {datetime.now().strftime('%Y-%m-%d')}"); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_page_break()

    doc.add_heading('测试用例说明', level=1)
    doc.add_paragraph('测试环境:')
    for x in ["操作系统: Windows 10/11", "Python 3.12+", "浏览器: Chrome/Firefox/Edge", "AI: DeepSeek API（需配置 DEEPSEEK_API_KEY）"]:
        doc.add_paragraph(x, style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('测试前准备:')
    for x in ["启动服务器: python app.py", "确认 http://localhost:5000 可访问", "确认 .env 中 DEEPSEEK_API_KEY 已配置", "准备测试文件: test_招标文件.docx"]:
        doc.add_paragraph(x, style='List Bullet')

    cases = [
        ("1.首页","TC-01","首页加载","服务器运行","访问 http://localhost:5000","页面正常显示，导航栏4个Tab，含上传区域","",""),
        ("1.首页","TC-02","上传表单检查","首页已加载","检查上传区域显示内容","拖拽上传提示文字正确","",""),
        ("1.首页","TC-03","导航切换","首页已加载","依次点击4个Tab","各页面加载正常，active高亮","",""),
        ("2.上传","TC-04","上传docx","文件就绪","选择 test_招标文件.docx 上传","跳转到处理进度页","",""),
        ("2.上传","TC-05","上传非docx","首页已加载","上传 .txt 文件","提示「请上传 .docx 格式」","",""),
        ("2.上传","TC-06","空上传","首页已加载","不选文件直接点上传","重定向回首页","",""),
        ("2.上传","TC-07","超大文件","首页已加载","上传超过50MB的文件","返回413错误","",""),
        ("3.解析","TC-08","进度显示","上传后跳转处理页","观察进度条和阶段指示器","进度条增长，阶段依次高亮","",""),
        ("3.解析","TC-09","解析完成","TC-08完成","等待进度100%→自动跳转","跳转到分析结果页，显示项目名和需求","",""),
        ("3.解析","TC-10","超时检测","处理页加载中","断开服务器连接","30秒后显示「服务器无响应」","",""),
        ("3.解析","TC-11","数据验证","分析完成","检查技术需求数量和项目名称","需求数>0，名称正确","",""),
        ("4.结果","TC-12","结果展示","分析完成","访问/analyze_result","显示需求列表和评分","",""),
        ("5.设计","TC-13","设计建议","分析完成","访问/design","显示技术方案建议","",""),
        ("6.大纲","TC-14","大纲生成","分析完成","访问/outline","章节数>0","",""),
        ("6.大纲","TC-15","大纲编辑","大纲已生成","修改章节标题并保存","修改成功","",""),
        ("7.生成","TC-16","启动生成","分析+大纲就绪","访问/generate，观察进度","进度条更新，15章60-90s完成","",""),
        ("7.生成","TC-17","生成结果","生成完成","检查方案内容","15个章节，标题和内容完整","",""),
        ("7.生成","TC-18","失败处理","改错API Key","重新上传生成","显示失败信息","",""),
        ("8.审核","TC-19","AI审核","方案已生成","访问/review","显示百分制评分和改进建议","",""),
        ("9.优化","TC-20","启动优化","审核完成","点击「开始优化」","优化进度页显示进度","",""),
        ("9.优化","TC-21","优化验证","优化完成","检查方案","版本号变为V1.1，内容已更新","",""),
        ("9.优化","TC-22","PDCA历史","≥1次优化","访问/pdca/history","显示版本号和评分变化","",""),
        ("10.下载","TC-23","下载docx","方案已生成","点击「下载方案」","文件名含版本号，>50KB，含里程碑表","",""),
        ("11.审批","TC-24","提交审批","方案已生成","填写信息点击提交","审批提交成功，跳转列表","",""),
        ("11.审批","TC-25","审批通过","有待审批","点击「通过」填意见","状态变为approved","",""),
        ("11.审批","TC-26","审批驳回","有待审批","点击「驳回」填原因","状态变为rejected","",""),
        ("12.发布","TC-27","审批后发布","审批已通过","访问/publish","显示「审批已通过」，有下载按钮","",""),
        ("13.参考","TC-28","上传参考","无","上传 test_招标文件.docx","显示风格摘要/结构/技术表述","",""),
        ("13.参考","TC-29","风格预览","有参考","点击「预览风格」","弹出风格示范文本","",""),
        ("13.参考","TC-30","删除参考","有参考","点击「删除」确认","从列表移除","",""),
        ("13.参考","TC-31","清空库","多篇参考","点击「清空文库」确认","全部删除","",""),
        ("14.仿写","TC-32","选择参考","有参考","访问/reference/mimic勾选文档","卡片显示选中状态","",""),
        ("14.仿写","TC-33","填写信息","已选参考","填项目名和需求","字段可正常填写","",""),
        ("14.仿写","TC-34","启动仿写","表单已填","点击「开始仿写生成」","跳转到进度页，实时显示完成数","",""),
        ("14.仿写","TC-35","动态大纲","选AI智能大纲","提交生成","章节数5-8章，标题与需求相关","",""),
        ("14.仿写","TC-36","仿写结果","仿写完成","检查结果页","Markdown渲染正确，章节可折叠","",""),
        ("14.仿写","TC-37","下载","仿写完成","点击下载","文件名格式: 仿写方案_ID.docx","",""),
        ("14.仿写","TC-38","版本历史","有多版本","访问/mimic/history","显示版本列表","",""),
        ("14.仿写","TC-39","版本对比","≥2版本","选择两个版本点击对比","左右分栏显示字数变化和相似度","",""),
        ("14.仿写","TC-40","版本恢复","有历史版本","点击「恢复」确认","当前方案被替换","",""),
        ("14.仿写","TC-41","反馈优化","仿写完成","勾选章节→填意见→提交","提示优化已启动","",""),
        ("15.Bot","TC-42","Bot页面","无","访问/bot","显示对话区域和命令列表","",""),
        ("15.Bot","TC-43","/status","Bot已加载","输入/status发送","返回项目状态","",""),
        ("15.Bot","TC-44","/help","Bot已加载","输入/help发送","显示命令列表","",""),
        ("15.Bot","TC-45","飞书webhook","无","POST /bot/webhook/feishu","返回challenge原值","",""),
        ("15.Bot","TC-46","企微webhook","无","POST /bot/webhook/wecom","返回reply","",""),
        ("16.全局","TC-47","重置","有活跃项目","点击「重新开始」确认","回到首页，storage清空","",""),
        ("16.全局","TC-48","响应式","无","缩放窗口到375px","自适应无横向滚动","",""),
        ("16.全局","TC-49","错误日志","无","遍历所有页面","控制台无Error日志","",""),
        ("16.全局","TC-50","线程稳定","生成进行中","生成中多次刷新页面","后台线程不受影响","",""),
        ("16.全局","TC-51","格式后处理","方案已生成","下载docx用Word打开","含里程碑表+SOP表","",""),
        ("16.全局","TC-52","会话状态","无","重启服务器后刷新","状态丢失，恢复初始","",""),
    ]

    current_module = ""
    for case in cases:
        module = case[0]
        if module != current_module:
            current_module = module
            doc.add_heading(f'模块 {module}', level=1)
        t = doc.add_table(rows=9, cols=2)
        t.style = 'Table Grid'
        r0 = t.rows[0]
        r0.cells[0].merge(r0.cells[1])
        r0.cells[0].text = f"{case[1]}: {case[2]}"
        for p in r0.cells[0].paragraphs:
            for rn in p.runs: rn.font.bold = True; rn.font.size = Pt(10)
        shd(r0.cells[0], 'E8F0FE')
        fields = [
            ("前置条件", case[3]), ("测试步骤", case[4]), ("预期结果", case[5]),
            ("实际结果", ""), ("状态", ""), ("测试人员", ""), ("测试日期", ""), ("备注", ""),
        ]
        for i, (l, v) in enumerate(fields, 1):
            t.rows[i].cells[0].text = l
            t.rows[i].cells[1].text = v
            for p in t.rows[i].cells[0].paragraphs:
                for rn in p.runs: rn.font.bold = True; rn.font.size = Pt(9)
            for p in t.rows[i].cells[1].paragraphs:
                for rn in p.runs: rn.font.size = Pt(9)
            t.rows[i].cells[0].width = Inches(1.0)
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading('测试结果汇总', level=1)
    st = doc.add_table(rows=8, cols=2)
    st.style = 'Table Grid'
    arow(st, ["模块", "用例数"], h=True)
    mods = {}
    for c in cases:
        m = c[0]
        mods[m] = mods.get(m, 0) + 1
    i = 1
    for m, cnt in mods.items():
        arow(st, [m, str(cnt)])
        i += 1
    arow(st, ["合计", str(len(cases))])

    p = doc.add_paragraph()
    p.add_run("\n通过率: ").bold = True; p.add_run("___ / " + str(len(cases)))
    p = doc.add_paragraph(); p.add_run("测试结论: ").bold = True
    doc.save(RPATH)
    print(f"测试用例已生成: {RPATH} ({len(cases)} 条)")

if __name__ == '__main__':
    build()
