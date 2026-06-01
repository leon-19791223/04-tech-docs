# -*- coding: utf-8 -*-
"""生成仿写功能优化方案 Word 文档."""
import os, sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装 python-docx")
    sys.exit(1)

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "Desktop", "投标系统", "仿写功能优化方案.docx")

def set_cell_shading(cell, color):
    shd = cell._element.get_or_add_tcPr().makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    cell._element.get_or_add_tcPr().append(shd)

def add_table_row(table, cells_data, header=False, color='2F5496'):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]; cell.text = str(text)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(9) if not header else Pt(10)
                run.font.bold = header
                if header: run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if header: set_cell_shading(cell, color)

def build_report():
    doc = Document()
    style = doc.styles['Normal']; style.font.name = '微软雅黑'; style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ===== Cover =====
    for _ in range(6): doc.add_paragraph()
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("投标系统 V3\n仿写功能优化方案"); run.font.size = Pt(28); run.font.bold = True; run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    doc.add_paragraph()
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"版本: V1.0 | 日期: {datetime.now().strftime('%Y-%m-%d')}"); run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_page_break()

    # ===== TOC =====
    doc.add_heading('目录', level=1)
    for item in ["1. 当前状态分析", "2. 功能问题清单", "3. 界面问题清单", "4. 整体优化方案", "5. 功能优化详情", "6. 界面优化详情", "7. 架构变更", "8. 实施计划", "9. 预期效果"]:
        p = doc.add_paragraph(item); p.paragraph_format.space_before = Pt(6); p.runs[0].font.size = Pt(12)
    doc.add_page_break()

    # ===== 1. Current State =====
    doc.add_heading('1. 当前状态分析', level=1)
    doc.add_paragraph('仿写功能自 V3 版本引入，允许用户上传已中标文档作为参考，AI 模仿其风格生成新方案。')
    doc.add_paragraph('当前技术架构:')

    arch = doc.add_table(rows=6, cols=2)
    arch.style = 'Table Grid'
    add_table_row(arch, ["组件", "现状"], header=True)
    arch_data = [
        ("参考上传", "拖拽上传 .docx → LLM 风格提取 → 存入内存 _library 列表"),
        ("风格提取", "LLM 分析语气/结构/技术表述三项特征，不可用时降级到通用模板"),
        ("仿写表单", "7 章硬编码大纲，无需求分析，仅传 project_name + background"),
        ("生成引擎", "复用 generator.generate_chapter()，追加 reference_context 到 prompt"),
        ("进度页", "前端轮询 + 模拟进度条"),
        ("结果页", "章节列表 ≤800 字预览 + docx 下载"),
    ]
    for row_data in arch_data:
        add_table_row(arch, row_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("代码规模: ").bold = True
    p.add_run("modules/reference.py (152 行) + 5 个模板 + app.py 约 130 行路由")
    doc.add_page_break()

    # ===== 2. Functional Issues =====
    doc.add_heading('2. 功能问题清单', level=1)

    issues_func = [
        ("F1", "高", "无需求分析",
         "仿写使用空的 tech_requirements/biz_requirements，不分析项目背景中的需求。"
         "生成方案内容空泛，缺乏针对性。"),
        ("F2", "高", "大纲不可配置",
         "7 章硬编码大纲（项目概述、总体方案、技术方案详解等），无法增删改章节，"
         "不适用不同场景（如纯技术方案 vs 商务方案）。"),
        ("F3", "中", "风格提取仅 3 项指标",
         "当前只提取语气、结构、技术表述三项，缺少句式样本、过渡词使用、"
         "标题风格、数据呈现方式等关键风格指标，模仿效果有限。"),
        ("F4", "中", "无风格预览",
         "用户在上传后看不到风格提取的示例效果，无法判断 AI 是否理解正确。"
         "需要提供风格模拟预览。"),
        ("F5", "中", "不支持章节级选择",
         "用户不能选择「参考 A 的总体风格 + 参考 B 的技术方案结构」，"
         "只能整篇参考所有文档的混合风格。"),
        ("F6", "低", "无多轮优化",
         "仿写结果不支持反馈优化，一次生成完成后无法针对章节进行迭代完善。"),
        ("F7", "低", "无版本比较",
         "多次仿写生成的结果无法对比，每次覆盖上一次的结果。"),
    ]

    t1 = doc.add_table(rows=len(issues_func)+1, cols=4)
    t1.style = 'Table Grid'
    add_table_row(t1, ["编号", "优先级", "问题", "说明"], header=True)
    for row_data in issues_func:
        add_table_row(t1, row_data)
    for row in t1.rows[1:]:
        row.cells[0].width = Inches(0.4)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(1.2)
        row.cells[3].width = Inches(3.8)
    doc.add_page_break()

    # ===== 3. UI Issues =====
    doc.add_heading('3. 界面问题清单', level=1)

    issues_ui = [
        ("U1", "高", "上传/仿写分离",
         "reference_upload.html（上传）和 reference_mimic.html（填写信息）是两页，"
         "用户上传后需跳转两次才能开始仿写。应合并为向导式流程。"),
        ("U2", "中", "结果页内容截断",
         "章节内容用 <pre> 标签显示，仅展示前 800 字符。无展开/折叠、"
         "无 Markdown 渲染（粗体、列表、表格不可见）。"),
        ("U3", "中", "进度页模拟进度",
         "进度条使用前端模拟 (setInterval)，非真实 LLM 进度。"
         "用户无法看到当前正在生成哪个章节的具体内容。"),
        ("U4", "中", "无参考文档对比",
         "结果页没有显示「本章节主要参考了哪个文档」，"
         "用户无法判断 AI 是否成功模仿了目标风格。"),
        ("U5", "低", "无编辑功能",
         "生成后的结果不支持在线编辑修改，用户只能下载 docx 后用 Word 修改。"),
        ("U6", "低", "移动端适配差",
         "表单和参考卡片在移动端布局错位，内联样式缺乏响应式支持。"),
    ]

    t2 = doc.add_table(rows=len(issues_ui)+1, cols=4)
    t2.style = 'Table Grid'
    add_table_row(t2, ["编号", "优先级", "问题", "说明"], header=True)
    for row_data in issues_ui:
        add_table_row(t2, row_data)
    for row in t2.rows[1:]:
        row.cells[0].width = Inches(0.4)
        row.cells[1].width = Inches(0.6)
        row.cells[2].width = Inches(1.2)
        row.cells[3].width = Inches(3.8)
    doc.add_page_break()

    # ===== 4. Overall Plan =====
    doc.add_heading('4. 整体优化方案', level=1)
    doc.add_paragraph('优化方案分为三个优先级阶段实施:')
    doc.add_paragraph()

    plan = doc.add_table(rows=4, cols=4)
    plan.style = 'Table Grid'
    add_table_row(plan, ["阶段", "时间", "目标", "涉及问题"], header=True)
    plan_data = [
        ("Phase 1\n快速优化", "1-2 天",
         "修复核心体验问题，提升可用性",
         "F1+F2+F3+U1+U2"),
        ("Phase 2\n体验升级", "3-5 天",
         "完整的向导式交互和实时预览",
         "F4+F5+U3+U4+U6"),
        ("Phase 3\n高级功能", "1 周",
         "多轮优化、版本管理、在线编辑",
         "F6+F7+U5"),
    ]
    for row_data in plan_data:
        add_table_row(plan, row_data)
    doc.add_page_break()

    # ===== 5. Functional Optimization =====
    doc.add_heading('5. 功能优化详情 (Phase 1+2)', level=1)

    doc.add_heading('5.1 新增仿写需求分析 (F1)', level=2)
    p = doc.add_paragraph()
    p.add_run("现状: ").bold = True
    p.add_run("仿写使用空 analysis，生成方案无需求针对性。")
    p = doc.add_paragraph()
    p.add_run("优化方案: ").bold = True
    p.add_run("在 reference_mimic_generate() 中新增轻量需求分析步骤:")
    p = doc.add_paragraph()
    p.add_run("1. ").bold = True
    p.add_run("用户在表单中可输入「核心需求」（文本框，每行一条）")
    p = doc.add_paragraph()
    p.add_run("2. ").bold = True
    p.add_run("AI 自动从 project_name + background + 核心需求 提取关键词作为 tech_requirements")
    p = doc.add_paragraph()
    p.add_run("3. ").bold = True
    p.add_run("大纲根据需求动态生成，而非硬编码 7 章")
    code = doc.add_paragraph()
    run = code.add_run(
        '# 优化后: 用户可输入需求\n'
        'requirements_text = request.form.get("requirements", "")\n'
        'reqs = [r.strip() for r in requirements_text.split("\\n") if r.strip()]\n'
        'analysis["tech_requirements"] = reqs or [f"{project_name} 整体建设需求"]\n\n'
        '# 大纲动态生成\n'
        'if reqs:\n'
        '    outline = generate_mimic_outline(analysis, reference_context)\n'
        'else:\n'
        '    outline = _get_default_outline()  # 7 章默认'
    )
    run.font.name = 'Consolas'; run.font.size = Pt(9)

    doc.add_heading('5.2 新增动态大纲生成 (F2)', level=2)
    p = doc.add_paragraph()
    p.add_run("现状: ").bold = True
    p.add_run("硬编码 7 章大纲，无法定制。")
    p = doc.add_paragraph()
    p.add_run("优化方案: ").bold = True
    p.add_run("提供三种大纲生成方式:")
    items = [
        "AI 生成: 传入 analysis + reference_context，LLM 自动生成匹配的章节结构",
        "模板选择: 提供「技术方案」「商务方案」「综合方案」三种预设大纲模板",
        "自定义: 用户在仿写表单中自由添加/删除/排序章节",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    p = doc.add_paragraph()
    p.add_run("新增函数: ").bold = True
    p.add_run("generate_mimic_outline(analysis, refs) → List[dict] 放在 modules/reference.py 中")

    doc.add_heading('5.3 增强风格提取 (F3)', level=2)
    p = doc.add_paragraph()
    p.add_run("现状: ").bold = True
    p.add_run("仅 3 项风格指标。")
    p = doc.add_paragraph()
    p.add_run("优化方案: ").bold = True
    p.add_run("增加至 6 项指标，包含示例文本:")
    doc.add_paragraph()

    style_table = doc.add_table(rows=7, cols=3)
    style_table.style = 'Table Grid'
    add_table_row(style_table, ["指标", "说明", "示例"], header=True)
    style_data = [
        ("style_summary", "总体风格描述", "语气正式严谨，偏好被动语态"),
        ("structure_outline", "段落/章节结构", "总-分-总结构，每章以引言开头"),
        ("tech_highlights", "技术表述特点", "偏好量化指标，常用「不低于」「≥」"),
        ("sentence_sample", "典型句式（新增）", "「本项目拟采用…」「基于…设计理念」"),
        ("transition_style", "过渡与连接词（新增）", "习惯用「此外」「与此同时」作为段落开头"),
        ("heading_style", "标题风格（新增）", "动词开头式标题：「建设…」「设计…」"),
    ]
    for row_data in style_data:
        add_table_row(style_table, row_data)

    doc.add_heading('5.4 风格预览功能 (F4)', level=2)
    p = doc.add_paragraph()
    p.add_run("优化方案: ").bold = True
    p.add_run("上传完成后，在参考库卡片中增加「生成风格示例」按钮。")
    doc.add_paragraph("点击后 LLM 基于提取的风格生成一段约 200 字的示范文本，用户确认风格符合预期后再开始仿写。", style='List Bullet')
    doc.add_paragraph("如不满意可重新提取或手动编辑风格描述。", style='List Bullet')
    doc.add_page_break()

    # ===== 6. UI Optimization =====
    doc.add_heading('6. 界面优化详情', level=1)

    doc.add_heading('6.1 向导式仿写流程 (U1)', level=2)
    p = doc.add_paragraph()
    p.add_run("现状: ").bold = True
    p.add_run("上传页 → 参考库 → 仿写表单，多页跳转。")
    p = doc.add_paragraph()
    p.add_run("优化方案: ").bold = True
    p.add_run("三步向导，单页面完成:")
    doc.add_paragraph()
    doc.add_paragraph("Step 1 — 上传参考: 拖拽上传 + 实时风格提取状态", style='List Bullet')
    doc.add_paragraph("Step 2 — 选择参考 + 填写信息: 参考卡片可多选，填入项目和需求", style='List Bullet')
    doc.add_paragraph("Step 3 — 确认生成: 摘要预览（参考数 + 大纲章节数 + 风格要点）", style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph("新增模板文件: templates/mimic_wizard.html（替代原有的 reference_upload.html + reference_mimic.html）")

    doc.add_heading('6.2 Markdown 渲染结果页 (U2)', level=2)
    p = doc.add_paragraph()
    p.add_run("现状: ").bold = True
    p.add_run("<pre> 标签截断显示，格式不可见。")
    p = doc.add_paragraph()
    p.add_run("优化方案: ").bold = True
    p.add_run("引入简单 JS Markdown 渲染器（marked.js 或自行实现），替换 <pre>:")
    doc.add_paragraph("章节内容可展开/折叠（默认折叠长章节）", style='List Bullet')
    doc.add_paragraph("Markdown 格式正确渲染（粗体、列表、代码块、表格）", style='List Bullet')
    doc.add_paragraph("章节锚点导航，点击左侧目录跳转到对应章节", style='List Bullet')
    code = doc.add_paragraph()
    run = code.add_run(
        '<!-- 使用简单的 Markdown 渲染 -->\n'
        '<div class="chapter-content markdown-body">\n'
        '  {{ chapter.content | markdown_render }}\n'
        '</div>\n\n'
        '# 新增 Jinja2 过滤器 (app.py)\n'
        '@app.template_filter("markdown_render")\n'
        'def markdown_render(text):\n'
        '    import re\n'
        '    text = re.sub(r"\\*\\*(.*?)\\*\\*", r"<b>\\1</b>", text)\n'
        '    text = re.sub(r"\\n- ", r"\\n<li>", text)\n'
        '    return text.replace("\\n", "<br>")'
    )
    run.font.name = 'Consolas'; run.font.size = Pt(9)

    doc.add_heading('6.3 实时章节进度 (U3)', level=2)
    p = doc.add_paragraph()
    p.add_run("现状: ").bold = True
    p.add_run("模拟进度条，不反映真实 LLM 进度。")
    p = doc.add_paragraph()
    p.add_run("优化方案: ").bold = True
    p.add_run("改为真实进度 + 实时内容流式展示:")
    doc.add_paragraph("每完成一章立即在页面追加该章标题和摘要（100 字），类似流式输出效果", style='List Bullet')
    doc.add_paragraph("进度条基于真实 completed/total 计算", style='List Bullet')
    doc.add_paragraph("显示当前章节的实时生成状态「AI 正在编写: 技术方案详解...」", style='List Bullet')

    doc.add_heading('6.4 参考来源标注 (U4)', level=2)
    p = doc.add_paragraph()
    p.add_run("优化方案: ").bold = True
    p.add_run("在结果页每章节旁标注「风格来源」:")
    doc.add_paragraph("如章节主要参考了某个文档的风格，显示该文档的文件名和风格摘要", style='List Bullet')
    doc.add_paragraph("多参考源时显示混合比例", style='List Bullet')

    doc.add_heading('6.5 结果页对比视图', level=1)
    p = doc.add_paragraph()
    p.add_run("新增功能: ").bold = True
    p.add_run("结果页增加「对比模式」，可同时查看仿写方案和参考文档风格样本:")
    doc.add_paragraph("左右分栏: 左栏为参考风格样本，右栏为 AI 生成内容", style='List Bullet')
    doc.add_paragraph("高亮显示风格匹配的关键句式", style='List Bullet')
    doc.add_page_break()

    # ===== 7. Architecture =====
    doc.add_heading('7. 架构变更', level=1)
    doc.add_paragraph('以下是对应优化需要的新增/修改文件清单:')

    arch_table = doc.add_table(rows=13, cols=3)
    arch_table.style = 'Table Grid'
    add_table_row(arch_table, ["文件", "变更类型", "变更内容"], header=True)
    arch_data = [
        ("modules/reference.py", "修改", "新增 generate_mimic_outline()、extract_style_via_llm_v2()"),
        ("modules/generator.py", "修改", "优化 reference_context 使用方式，支持章节级风格指定"),
        ("app.py", "修改", "新增 /mimic/wizard 路由、mimic_outline 接口、markdown_render 过滤器"),
        ("templates/mimic_wizard.html", "新增", "三步向导式仿写页面"),
        ("templates/mimic_result_v2.html", "新增", "新版结果页（Markdown 渲染 + 对比视图）"),
        ("templates/mimic_generating_v2.html", "新增", "新版进度页（实时追加章节）"),
        ("templates/reference_upload.html", "修改", "增加风格预览和重新提取功能"),
        ("templates/reference_mimic.html", "修改", "增加需求输入框、大纲模板选择、自定义章节"),
        ("templates/reference_mimic_result.html", "修改", "Markdown 渲染 + 可展开折叠 + 对比模式入口"),
        ("static/style.css", "修改", "新增向导步骤条样式、Markdown 正文样式"),
        ("static/mimic.js", "新增", "仿写相关前端交互逻辑（向导、展开折叠、对比切换）"),
        ("test_mimic_features.py", "新增", "仿写功能专项测试"),
    ]
    for row_data in arch_data:
        add_table_row(arch_table, row_data)

    doc.add_paragraph()
    doc.add_paragraph('核心数据流变化:')
    doc.add_paragraph('用户输入需求 → AI 生成动态大纲 → 参考风格注入 prompt → 逐章生成 → Markdown 渲染', style='List Bullet')
    doc.add_page_break()

    # ===== 8. Implementation Plan =====
    doc.add_heading('8. 实施计划', level=1)

    doc.add_heading('Phase 1: 快速优化 (1-2 天)', level=2)
    p1 = [
        ("Day 1", [
            "F1: 仿写表单增加需求输入框，analysis 集成用户需求",
            "F3: 风格提取由 3 项扩至 6 项指标",
            "U1: 上传页/仿写页合并，减少跳转步骤",
            "U4: 结果页增加参考来源标注",
        ]),
        ("Day 2", [
            "F2: 实现 generate_mimic_outline() LLM 动态生成大纲",
            "U2: Markdown 渲染过滤器 + 展开折叠",
            "U3: 进度页改为真实进度显示",
            "测试: 跑全量测试 + 仿写专项测试",
        ]),
    ]
    for day, tasks in p1:
        p = doc.add_paragraph()
        p.add_run(day + ": ").bold = True
        p.add_run(" | ".join(tasks))

    doc.add_heading('Phase 2: 体验升级 (3-5 天)', level=2)
    p2 = [
        ("Day 3-4", [
            "F4: 风格预览功能（上传后生成示范文本）",
            "F5: 章节级风格选择（选择 A 的某章 + B 的某章）",
            "U5: 结果页在线编辑（contenteditable）",
        ]),
        ("Day 5", [
            "U6: 移动端响应式适配",
            "向导式 UI 完善",
            "对比视图功能",
            "全量测试回归",
        ]),
    ]
    for day, tasks in p2:
        p = doc.add_paragraph()
        p.add_run(day + ": ").bold = True
        p.add_run(" | ".join(tasks))

    doc.add_heading('Phase 3: 高级功能 (1 周)', level=2)
    p3 = [
        ("Week 2", [
            "F6: 仿写结果多轮优化（章节级反馈 → 重新生成）",
            "F7: 版本管理（历史仿写记录保存 + 对比）",
            "U5: 在线编辑器完善 + 自动保存",
            "AI 风格匹配度评分",
        ]),
    ]
    for day, tasks in p3:
        p = doc.add_paragraph()
        p.add_run(day + ": ").bold = True
        p.add_run(" | ".join(tasks))

    doc.add_page_break()

    # ===== 9. Expected Results =====
    doc.add_heading('9. 预期效果', level=1)

    exp_table = doc.add_table(rows=8, cols=3)
    exp_table.style = 'Table Grid'
    add_table_row(exp_table, ["维度", "优化前", "优化后"], header=True)
    exp_data = [
        ("参考上传到生成", "3 步（上传→库→表单→生成）", "1 步（向导式，一站式完成）"),
        ("大纲定制", "7 章固定大纲", "AI 动态生成 + 3 种模板 + 自定义"),
        ("需求针对性", "无需求分析，内容空泛", "用户输入需求 → 针对性生成"),
        ("风格提取", "3 项通用指标", "6 项精细化指标 + 示例文本"),
        ("结果查看", "<pre> 截断 800 字", "Markdown 渲染 + 展开折叠 + 导航"),
        ("风格反馈", "无", "风格预览 + 参考来源标注 + 对比视图"),
        ("迭代优化", "不支持", "章节级反馈 → 重新生成"),
    ]
    for row_data in exp_data:
        add_table_row(exp_table, row_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("总结: ")
    run.font.bold = True
    p.add_run(
        "本次优化将使仿写功能从「基础的风格模仿」升级为「需求驱动的智能仿写」，"
        "用户可灵活定制大纲和风格粒度，生成结果可视化程度大幅提升。"
        "Phase 1 解决核心痛点，Phase 2+3 构建完整体验闭环。"
    )

    doc.save(REPORT_PATH)
    return REPORT_PATH

if __name__ == '__main__':
    path = build_report()
    print(f"方案已生成: {path}")
