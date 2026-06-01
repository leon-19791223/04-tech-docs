# -*- coding: utf-8 -*-
"""生成 DWS部署交互助手-开发功能建议书.docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def make_header_row(table, row_idx, texts):
    """Style a header row with blue background"""
    for i, text in enumerate(texts):
        cell = table.cell(row_idx, i)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, '1A3C6E')

def add_data_row(table, row_idx, texts, bold_first=False):
    """Add a data row with optional cell shading"""
    for i, text in enumerate(texts):
        cell = table.cell(row_idx, i)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if bold_first and i == 0:
            run.bold = True
        # Alternating row color
        if row_idx % 2 == 0:
            set_cell_shading(cell, 'F0F4F8')

def add_styled_table(doc, headers, data, col_widths=None):
    """Create a styled table"""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    make_header_row(table, 0, headers)
    # Data
    for ri, row in enumerate(data):
        add_data_row(table, ri + 1, row, bold_first=True)
    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table

def add_separator(doc):
    doc.add_paragraph('─' * 50)

# ════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DWS 部署交互助手')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('开发功能建议书')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基于 DWS 部署 V2.0 文档的交互式部署工具设计')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026 年 5 月 30 日')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ════════════════════════════════════════════
# 目录页（手动）
# ════════════════════════════════════════════
doc.add_heading('目  录', level=1)
toc_items = [
    '第一章  部署流程分析 ........................................ 3',
    '第二章  交互助手定位与架构 .................................. 5',
    '第三章  核心功能设计 ........................................ 6',
    '第四章  技术选型 ............................................ 9',
    '第五章  开发路线 ........................................... 10',
    '第六章  与现有工具对比 ..................................... 11',
    '第七章  待补充项与优化建议 .................................. 12',
    '第八章  博客案例对比与补充 .................................. 17',
    '第九章  网上资源综合分析与建议 .............................. 21',
    '第十章  DWS 运维管理补充 ..................................... 25',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.line_spacing = 2.0

doc.add_page_break()

# ════════════════════════════════════════════
# 第一章
# ════════════════════════════════════════════
doc.add_heading('第一章  部署流程分析', level=1)

doc.add_paragraph(
    '华为 FusionInsight DWS 9.1.0 在 KylinOS ARM64 环境下的完整部署流程涉及 40+ 个操作步骤，'
    '可归纳为以下 8 大阶段：'
)

# 8大阶段表格
headers = ['阶段', '内容', '步骤数']
data = [
    ['阶段1：环境准备', '/etc/hosts, hostname, Python 3.8.5, yum源(ISO), chrony/NTP', '5'],
    ['阶段2：OS调优', '关闭audit/firewall, 关闭透明大页, NUMA关闭, OOM配置,\nIO调度器, vm.min_free_kbytes, SELinux/swap/IPv6', '8'],
    ['阶段3：硬件验证', 'Hi1822网卡, RAID配置, SSD/HDD检测, FIO压测, dd I/O测试', '5'],
    ['阶段4：磁盘规划', 'parted分区, mkfs.xfs, UUID挂载, fstab', '4'],
    ['阶段5：网络验证', 'check_net, speed_test send/recv, sar/gsar, 带宽>800MB/s', '4'],
    ['阶段6：软件部署', '解压GaussDB包, FusionInsight_Manager, FusionInsight_SetupTool, cp packs', '4'],
    ['阶段7：预检查', 'preinstall.ini, setuptool.sh preinstall, precheck, rpm安装', '4'],
    ['阶段8：部署OMS', 'install.sh, 安装manager, 完成验证', '3'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')

# 部署全景图（文字）
doc.add_heading('1.1  部署全景流程', level=2)
panorama = (
    '环境准备 → OS调优 → 硬件验证 → 磁盘规划 → 网络验证 → 软件部署 → 预检查 → 部署OMS\n\n'
    '每个阶段有严格的依赖关系：\n'
    '• 环境准备是基础，必须先完成\n'
    '• OS调优和硬件验证可并行执行\n'
    '• 磁盘规划必须在软件部署前完成\n'
    '• 预检查通过后才能执行部署OMS'
)
p = doc.add_paragraph(panorama)
p.paragraph_format.line_spacing = 1.5

# 痛点分析
doc.add_heading('1.2  当前部署痛点', level=2)

headers = ['痛点', '说明']
data = [
    ['步骤多易遗漏', '40+ 步纯手动操作，序号跳转，容易漏步'],
    ['参数靠手记', 'preinstall.ini 的 IP/主机名/分区配置全靠手动编辑，易出错'],
    ['无顺序约束', '依赖关系不明确，前置未完成可能导致后续失败'],
    ['命令分散', 'FIO/dd/speed_test 等工具散落在文档中，每次要翻找'],
    ['结果不可追溯', '每个步骤执行成功/失败没有记录，排障靠回忆'],
    ['无模板化', '每个新环境都要重新编辑所有配置，重复劳动'],
]
add_styled_table(doc, headers, data)

doc.add_page_break()

# ════════════════════════════════════════════
# 第二章
# ════════════════════════════════════════════
doc.add_heading('第二章  交互助手定位与架构', level=1)

doc.add_heading('2.1  产品定位', level=2)
p = doc.add_paragraph(
    '一个 TUI（终端交互式）命令行工具，采用类似问答系统的交互风格，'
    '覆盖 DWS 部署全流程的向导式交互。'
)
p.paragraph_format.line_spacing = 1.5

p2 = doc.add_paragraph(
    '核心设计理念：让部署工程师从"翻阅文档+复制命令"转变为"跟随向导+确认执行"，'
    '降低人为失误，提升部署效率与可追溯性。'
)
p2.paragraph_format.line_spacing = 1.5

doc.add_heading('2.2  系统架构', level=2)

# 架构图用等宽字体的段落
arch_text = (
    'dws_deploy_assistant.py\n'
    '├── 核心引擎\n'
    '│   ├── 阶段管理器 (Phase Manager) — 8阶段顺序编排\n'
    '│   ├── 步骤引擎 (Step Engine) — 每阶段N步，带依赖检查\n'
    '│   ├── 配置中心 (Config Center) — 统一管理IP/主机名/分区参数\n'
    '│   └── 审计日志 (Audit Log) — 每个步骤结果记录到JSON\n'
    '│\n'
    '├── 交互层 (TUI)\n'
    '│   ├── 主菜单 — 展示8大阶段进度标记\n'
    '│   ├── 步骤执行 — 显示命令、执行、捕获输出、确认结果\n'
    '│   ├── 配置编辑 — 交互式编辑preinstall.ini参数\n'
    '│   └── 结果看板 — 每步成功/失败状态\n'
    '│\n'
    '└── 工具库\n'
    '    ├── 远程执行 (paramiko SSH) — 支持多节点并行\n'
    '    ├── 模板引擎 — fstab/preinstall.ini/grub模板化\n'
    '    └── 检查器 — 每个步骤前置条件自动校验'
)
p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)
p.paragraph_format.line_spacing = 1.2

doc.add_page_break()

# ════════════════════════════════════════════
# 第三章
# ════════════════════════════════════════════
doc.add_heading('第三章  核心功能设计', level=1)

doc.add_heading('3.1  配置中心', level=2)
doc.add_paragraph(
    '用户只需一次录入所有部署参数，后续所有配置文件自动生成，彻底消除手动编辑错误。'
)
doc.add_paragraph('核心配置项：')

headers = ['配置项', '说明', '示例']
data = [
    ['节点列表', '各节点管理IP/业务IP', '10.134.21.190~192'],
    ['主机名', '各节点主机名', 'uatdws0001~0003'],
    ['OS版本', '操作系统类型及版本', 'Kylin-V10-SP2 ARM64'],
    ['内存大小', '各节点物理内存', '256GB / 512GB'],
    ['数据盘', '数据盘设备名', '/dev/sdb, /dev/sdc'],
    ['ISO路径', 'KylinOS ISO镜像路径', '/opt/Kylin-...iso'],
    ['GaussDB包', 'GaussDB压缩包路径', '/opt/GaussDB_MPPDB_9.1.0.zip'],
]
add_styled_table(doc, headers, data)

doc.add_heading('3.2  步骤执行引擎', level=2)
doc.add_paragraph('每个步骤的执行流程：')
steps = [
    '显示将要执行的命令及说明',
    '用户确认执行（或选择跳过）',
    '自动执行命令并实时捕获输出',
    '根据返回码和输出内容自动判断成功/失败',
    '记录执行结果到审计日志',
    '自动推进到下一个步骤（或停留在失败步骤等待处理）',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('3.3  智能模板生成', level=2)
doc.add_paragraph('基于配置中心的数据，一键生成以下部署配置文本：')
templates = [
    '/etc/hosts — 所有节点主机映射',
    'preinstall.ini — FusionInsight 完整预安装配置',
    '/etc/fstab — 磁盘分区自动挂载配置',
    'grub 参数 — 透明大页/IO调度器等内核参数',
    '分区脚本 — 批量 parted + mkfs.xfs 脚本',
]
for t in templates:
    doc.add_paragraph(f'• {t}')

doc.add_heading('3.4  审计日志', level=2)
doc.add_paragraph(
    '每次部署生成的结构化审计日志，包含完整的执行轨迹，'
    '支持导出为部署报告。'
)

headers = ['字段', '说明', '示例']
data = [
    ['session', '部署会话ID', '2026-05-30_1425'],
    ['cluster', '集群名称', 'uatdws'],
    ['phase', '阶段编号', '2'],
    ['step', '步骤标识', 'disable_firewall'],
    ['status', '执行状态', 'passed / failed'],
    ['cmd', '执行的命令', 'systemctl stop firewalld'],
    ['output', '命令输出', '（空）'],
    ['duration_sec', '耗时（秒）', '1.2'],
    ['timestamp', '时间戳', '2026-05-30T14:26:00'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')
p = doc.add_paragraph(
    '审计日志的价值：部署完成后可直接导出为部署报告交付客户，'
    '也可以用于回溯排查部署中的异常步骤。'
)
p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ════════════════════════════════════════════
# 第四章
# ════════════════════════════════════════════
doc.add_heading('第四章  技术选型', level=1)

doc.add_paragraph(
    '技术选型原则：与 DWS 目标环境保持一致，最小化额外依赖，确保在离线环境也能可用。'
)

headers = ['层', '技术', '理由']
data = [
    ['语言', 'Python 3.8+', '与 DWS 环境 Python 版本一致，无需额外安装'],
    ['终端UI', 'rich / textual', '终端彩色界面，进度条，表格展示，专业感强'],
    ['远程执行', 'paramiko', '多节点 SSH 远程命令执行，支持并发'],
    ['模板引擎', 'jinja2', '生成 fstab / preinstall.ini / grub 等配置'],
    ['日志存储', 'JSON Lines', '结构化可追溯，可解析导出报告'],
    ['分发方式', '单文件 + pip', 'scp 到任一节点即可运行，无需复杂部署'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')
p = doc.add_paragraph(
    '注：以上依赖均可通过离线 pip 包方式提前准备，适用于客户现场无互联网的环境。'
)
p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ════════════════════════════════════════════
# 第五章
# ════════════════════════════════════════════
doc.add_heading('第五章  开发路线', level=1)

doc.add_paragraph('按增量迭代方式分三个阶段交付，每个阶段解决核心痛点：')

doc.add_heading('5.1  v0.1 — 配置中心 + 模板生成', level=2)
doc.add_paragraph('目标：解决参数配置痛点')
items = [
    'IP/主机名/磁盘等参数交互式录入',
    '一键生成 preinstall.ini / fstab / hosts',
    '导出配置文件到指定路径',
    '预计开发周期：2-3天',
]
for item in items:
    doc.add_paragraph(f'• {item}')

doc.add_heading('5.2  v0.2 — 步骤执行引擎', level=2)
doc.add_paragraph('目标：解决步骤遗漏痛点')
items = [
    '8阶段顺序编排，带依赖检查',
    '本地及远程命令执行',
    '成功/失败自动标记',
    '审计日志记录',
    '预计开发周期：3-5天',
]
for item in items:
    doc.add_paragraph(f'• {item}')

doc.add_heading('5.3  v0.3 — 检查器 + 验证', level=2)
doc.add_paragraph('目标：解决结果不可追溯痛点')
items = [
    '前置条件自动检查（如磁盘是否存在、端口是否占用）',
    '结果自动判定（基于输出内容）',
    '部署报告导出（PDF/CSV）',
    '断点续做（从失败步骤继续）',
    '预计开发周期：3-5天',
]
for item in items:
    doc.add_paragraph(f'• {item}')

doc.add_page_break()

# ════════════════════════════════════════════
# 第六章
# ════════════════════════════════════════════
doc.add_heading('第六章  与现有工具对比', level=1)

doc.add_paragraph(
    '与当前纯手动部署方式及华为官方 FusionInsight SetupTool 的对比分析：'
)

headers = ['维度', '当前纯手动', 'FusionInsight SetupTool', '本工具']
data = [
    ['覆盖范围', '全部 40+ 步，靠文档', 'preinstall + precheck 部分', '全流程 8 阶段向导'],
    ['交互性', '翻阅文档 + 复制命令', '无交互，命令行参数', '交互式确认，每步确认'],
    ['配置管理', '手动编辑，极易出错', '需手动编辑 ini 文件', '一次录入，一键生成'],
    ['审计追踪', '无记录', '有日志但分散在各节点', '结构化审计日志'],
    ['断点续做', '无，失败从头来', '部分支持', '完整支持断点续做'],
    ['部署报告', '无', '无', '可导出 PDF/CSV 报告'],
]
add_styled_table(doc, headers, data)

doc.add_page_break()

# ════════════════════════════════════════════
# 第七章  待补充项与优化建议
# ════════════════════════════════════════════
doc.add_heading('第七章  待补充项与优化建议', level=1)

doc.add_paragraph(
    '对上述方案进行评审后，识别出以下 10 项待补充或优化的关键能力，'
    '按优先级从高到低排列。'
)

# 7.1 回滚机制
doc.add_heading('7.1  回滚机制（P0）', level=2)
doc.add_paragraph('当前方案只支持"断点续做"，缺少破坏性操作的回滚能力。实际部署中，以下步骤一旦出错需要立即恢复：')

headers = ['步骤', '风险', '回滚方式']
data = [
    ['parted 分区', '数据丢失', '保存原始分区表，提供恢复脚本'],
    ['grub2-mkconfig', '系统无法启动', '备份原 grub.cfg，支持一键恢复'],
    ['sysctl -p', '内核参数异常', '备份原 sysctl.conf，支持回退'],
    ['swapoff', '依赖检查', '记录原 swap 设备，支持重新启用'],
    ['SELinux 关闭', '安全策略变更', '备份原 config'],
]
add_styled_table(doc, headers, data)

# 7.2 前置全量检查
doc.add_heading('7.2  前置全量检查（P0）', level=2)
doc.add_paragraph(
    '在开始任何操作之前进行一次性全量环境检测，避免中途发现环境不达标。'
)

headers = ['检测项', '检测内容', '合格标准']
data = [
    ['OS 版本', 'cat /etc/os-release', 'Kylin-V10-SP2 及以上'],
    ['CPU 架构', 'uname -m', 'aarch64'],
    ['内存容量', 'free -g', '≥ 256GB'],
    ['数据盘数量', 'lsblk', '≥ 2 块独立数据盘'],
    ['网络连通性', 'ping 跨节点 IP', '全部互通'],
    ['Python 版本', 'python3 --version', '≥ 3.8'],
    ['ISO 镜像', 'ls -l ISO路径', '文件存在且完整'],
    ['GaussDB 包', 'md5sum', 'MD5 校验一致'],
]
add_styled_table(doc, headers, data)
doc.add_paragraph('')
doc.add_paragraph('检测结果用三色标识：🟢 绿色通过 / 🟡 黄色警告 / 🔴 红色不通过')

# 7.3 环境模板
doc.add_heading('7.3  环境模板复用（P1）', level=2)
doc.add_paragraph(
    '当前配置中心为单次录入，缺少模板复用能力。实际环境中通常有多套规格：'
)
templates = [
    '开发环境：3节点 + 256GB → 保存为 dev-template',
    '测试环境：3节点 + 512GB → 保存为 test-template',
    '生产环境：6节点 + 512GB → 保存为 prod-template',
    '下次部署同规格 → 加载模板 → 修改 IP → 直接部署',
]
for t in templates:
    doc.add_paragraph(f'• {t}')

# 7.4 基线比对
doc.add_heading('7.4  基线比对阈值判定（P1）', level=2)
doc.add_paragraph(
    '当前成功/失败判断仅依赖返回码，但 DWS 部署的很多步骤需要阈值判定：'
)

headers = ['步骤', '检测指标', '合格阈值', '当前判断', '改进后']
data = [
    ['FIO 写入', '带宽 MB/s', '≥ 800MB/s (SSD)', 'return code=0', '带宽≥800?'],
    ['FIO 读取', '带宽 MB/s', '≥ 1500MB/s', 'return code=0', '带宽≥1500?'],
    ['网络测试', '吞吐 MB/s', '≥ 800MB/s (10GE)', 'return code=0', '吞吐≥800?'],
    ['网络丢包', 'drop/resend %', '< 0.01%', '无检测', '自动解析gsar'],
    ['dd 写入', '速率 MB/s', '匹配磁盘规格', 'return code=0', '速率达标?'],
    ['内存', 'free -h', '≥ 256GB', '无检测', '自动解析'],
]
add_styled_table(doc, headers, data)

# 7.5 幂等性
doc.add_heading('7.5  幂等性设计（P1）', level=2)
doc.add_paragraph(
    '如果一个步骤已经执行过，再次执行时应自动检测已完状态并跳过，而不是重复执行：'
)
examples = [
    'systemctl stop firewalld → 第一次执行关闭；第二次检测到已关闭 → 跳过并显示 "✓ 已完成"',
    'sysctl -p → 检测参数是否已生效 → 已生效则跳过',
    'parted 分区 → 检测分区是否存在 → 存在则跳过',
    'mount -a → 检测挂载点是否已挂载 → 已挂载则跳过',
]
for ex in examples:
    doc.add_paragraph(f'• {ex}')

# 7.6 凭据管理
doc.add_heading('7.6  凭据管理（P2）', level=2)
doc.add_paragraph('部署过程涉及 SSH 多节点操作，需明确凭据管理策略：')

headers = ['方案', '说明', '推荐度']
data = [
    ['SSH 密钥对', '配置免密登录，适合自动化', '⭐⭐⭐ 推荐'],
    ['交互式输入', '每次部署输入密码，适合临时使用', '⭐⭐ 备选'],
    ['凭据文件加密', '存储加密的凭据文件，适合无人值守', '⭐ 谨慎使用'],
]
add_styled_table(doc, headers, data)

# 7.7 部署报告模板
doc.add_heading('7.7  部署报告模板（P2）', level=2)
doc.add_paragraph('审计日志有了，但缺少面向客户交付的部署报告。建议规范报告内容：')

headers = ['报告章节', '内容']
data = [
    ['集群基本信息', '节点列表、版本号、硬件规格'],
    ['步骤执行结果表', '每步骤通过/失败/跳过状态及耗时'],
    ['性能基线数据', 'FIO 测试结果、网络带宽、内存信息'],
    ['异常处理记录', '执行过程中的异常及处理方式'],
    ['部署结论', '通过 / 有条件通过 / 不通过'],
]
add_styled_table(doc, headers, data)
doc.add_paragraph('')
doc.add_paragraph('输出格式：PDF（正式交付）/ HTML（在线查看）')

# 7.8 并发执行策略
doc.add_heading('7.8  并发执行策略（P2）', level=2)
doc.add_paragraph(
    '多节点部署时，部分操作可以跨节点并行执行，显著缩短总耗时：'
)
examples = [
    '阶段2 OS调优 → N个节点同时执行相同命令 → 总耗时 = 单节点耗时',
    '阶段3 硬件验证 → 同时在各节点跑 FIO → 节点间无依赖',
    '阶段4 磁盘规划 → 同时在各节点分区 → 操作独立',
    '注意事项：并发数可配置（默认3），避免 SSH 连接风暴',
]
for ex in examples:
    doc.add_paragraph(f'• {ex}')

# 7.9 多版本支持
doc.add_heading('7.9  多版本与多OS支持（P3）', level=2)
doc.add_paragraph(
    '当前方案仅基于 DWS 9.1.0 + KylinOS ARM64，实际环境存在多种组合：'
)
versions = [
    'DWS 8.x → 步骤不同，cstore_buffers 等参数不同',
    'DWS 9.0.x → preinstall.ini 字段略有差异',
    '欧拉OS (EulerOS) → yum 源配置、ISO 挂载方式不同',
    '建议：步骤定义参数化/版本化，运行时根据检测到的版本自动适配',
]
for v in versions:
    doc.add_paragraph(f'• {v}')

# 7.10 单机/集群模式
doc.add_heading('7.10  单机 vs 集群模式区分（P3）', level=2)
doc.add_paragraph(
    '不同部署场景的目标和步骤完全不同，工具应自动识别：'
)

headers = ['模式', '适用场景', '步骤差异']
data = [
    ['单机模式', '学习测试、POC 验证', '无需 preinstall.ini / kerbose / 节点间 SSH'],
    ['集群模式', '生产部署、正式环境', '完整 8 阶段，多节点协同'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')
doc.add_paragraph('工具启动时根据检测到的节点数量自动判断模式（单节点→单机模式，多节点→集群模式）')

doc.add_paragraph('')

# 补充优先级总表
doc.add_heading('7.11  补充优先级汇总', level=2)

headers = ['优先级', '待补充项', '影响', '建议纳入版本']
data = [
    ['P0', '回滚机制', '破坏性操作无挽回途径', 'v0.2'],
    ['P0', '前置全量检查', '避免中途发现环境不达标', 'v0.1'],
    ['P1', '基线比对（阈值判定）', 'FIO/网络等结果误判', 'v0.2'],
    ['P1', '幂等性设计', '重复执行导致异常', 'v0.2'],
    ['P1', '环境模板复用', '多环境重复录入', 'v0.1'],
    ['P2', '凭据管理', 'SSH 认证方式不明确', 'v0.2'],
    ['P2', '部署报告模板', '客户交付需要正式报告', 'v0.3'],
    ['P2', '并发执行策略', '3节点串行部署耗时长', 'v0.2'],
    ['P3', '多版本/多OS支持', 'DWS不同版本参数差异', 'v1.0'],
    ['P3', '单机/集群模式区分', '学习和生产场景区分', 'v0.2'],
]
add_styled_table(doc, headers, data)

doc.add_page_break()

# ════════════════════════════════════════════
# 第八章  博客案例对比与补充
# ════════════════════════════════════════════
doc.add_heading('第八章  博客案例对比与补充', level=1)

doc.add_paragraph(
    '参考博客《麒麟V10 SP2安装GaussDB(DWS)-ESL（线下纯软版）》（2024-09，DWS 8.2.1），'
    '与本文档分析的 DWS 9.1.0 V2.0 部署文档进行对比，发现以下差异和补充点。'
)

doc.add_heading('8.1  基础信息对比', level=2)

headers = ['对比项', '本文分析（DWS V2.0）', '博客案例（DWS 8.2.1 ESL）']
data = [
    ['DWS 版本', '9.1.0', '8.2.1'],
    ['操作系统', 'KylinOS ARM64', 'KylinOS x86_64'],
    ['CPU 架构', 'aarch64', 'x86_64'],
    ['节点数量', '3节点（管理+数据合一）', '3节点（MN&CN&DN合一）'],
    ['磁盘配置', '元数据盘2 + 数据盘2', '管控节点4盘 / 数据节点2盘'],
    ['部署工具', 'FusionInsight_SetupTool', 'FusionInsight_SetupTool + LLD工具'],
]
add_styled_table(doc, headers, data)

doc.add_heading('8.2  新增发现的关键差异', level=2)

doc.add_heading('补充1：LLD 配置规划工具（新增前置阶段）', level=3)
doc.add_paragraph(
    '博客指出整个部署的起点是华为提供的 LLD（Low Level Design）Excel 配置规划工具，'
    '该工具生成所有后续需要的配置文件。我们的方案直接从 preinstall.ini 开始，跳过了这个关键环节。'
)
doc.add_paragraph('LLD 工具生成的文件清单：')
lld_files = [
    'preinstall.ini -- 预安装主配置',
    'host*.ini -- 各节点分区配置',
    'checkNodes.Config -- 预检查配置文件',
    'password.ini -- 节点密码文件',
    'install_oms/<IP>.ini -- 主备 Manager 安装配置',
    '*.xml -- Web 导入创建集群的模板',
]
for f in lld_files:
    doc.add_paragraph(f'  * {f}')
p = doc.add_paragraph()
run = p.add_run('建议：将 LLD 工具纳入交互助手流程，或提供配置生成器替代 LLD Excel。')
run.bold = True

doc.add_heading('补充2：磁盘预读块参数（read_ahead_kb）', level=3)
doc.add_paragraph(
    '所有节点必须设置磁盘预读块大小为 16384，否则 preinstall 环境检查会报 ERROR。'
)
for c in ['echo 16384 > /sys/block/sda/queue/read_ahead_kb',
          'echo 16384 > /sys/block/sdb/queue/read_ahead_kb',
          '写入 /etc/rc.local 实现重启持久化']:
    doc.add_paragraph(f'  * {c}')

doc.add_heading('补充3：KylinOS audit RPM 补丁', level=3)
doc.add_paragraph('麒麟操作系统存在已知 bug，需手动替换 audit 相关 RPM 包：')
for c in ['rpm -Uvh audit-3.0-5.se.12.ky10.x86_64.rpm',
          'rpm -Uvh audit-libs-3.0-5.se.12.ky10.x86_64.rpm',
          'rpm -Uvh python3-audit-3.0-5.se.12.ky10.x86_64.rpm',
          'systemctl daemon-reload && systemctl restart auditd.service']:
    doc.add_paragraph(f'  * {c}')

doc.add_heading('补充4：SSH 超时防护配置', level=3)
doc.add_paragraph('多节点部署耗时长，必须提前配置 SSH 防超时：')
for c in ['export TMOUT=0 写入 /etc/profile',
          'vi /etc/ssh/sshd_config -> ClientAliveInterval 0, ClientAliveCountMax 3',
          'service sshd restart']:
    doc.add_paragraph(f'  * {c}')

doc.add_heading('补充5：Manager 主备安装交互确认', level=3)
doc.add_paragraph(
    '安装备 Manager 节点时，工具会交互式确认浮动 IP 是否已存在。'
    '这需要交互助手支持非静默模式的交互响应。'
)

doc.add_heading('补充6：sudo 脚本更新步骤', level=3)
doc.add_paragraph('Manager 安装完成后、Web 创建集群前，必须执行 sudo 脚本更新，我们的方案完全遗漏。')
for c in ['cd /opt/FusionInsight_SetupTool/os_optimization_tool',
          'sh optimization_patch.sh install',
          'sh optimization_patch.sh commit']:
    doc.add_paragraph(f'  * {c}')

doc.add_heading('补充7：Web UI 创建集群（非 CLI）', level=3)
doc.add_paragraph('Manager 安装完成后，集群创建通过 Web 界面完成。操作流程：')
for s in ['https://<浮动IP>:28443/web',
          '创建集群 -> 模板安装 -> 上传 XML',
          '输入 root 密码发现主机 -> 设置机架',
          '选择服务 -> 分配角色 -> 提交安装']:
    doc.add_paragraph(f'  * {s}')

doc.add_heading('补充8：部署后验证清单', level=3)

headers = ['验证项', '验证方式', '合格标准']
data = [
    ['服务状态', 'Web UI -> 集群 -> 服务', '运行状态为"良好"'],
    ['节点状态', 'Web UI -> 主机', '运行状态为"良好"'],
    ['健康检查', 'Web UI 执行集群健康检查', '全部通过'],
    ['gsql 登录', 'gsql -d postgres -p25308 -r', '登录成功'],
    ['VIP 查询', 'cm_ctl query -Cvipd', 'VIP 正常'],
]
add_styled_table(doc, headers, data)

doc.add_heading('8.3  对本方案的更新建议', level=2)

headers = ['新增/修改项', '类型', '影响版本']
data = [
    ['增加 LLD 工具阶段（阶段0）', '新增前置阶段', 'v0.1'],
    ['read_ahead_kb 配置检测', '补充 OS 调优', 'v0.2'],
    ['KylinOS audit 补丁检测', '补充环境检查', 'v0.1'],
    ['SSH 超时配置', '补充环境准备', 'v0.1'],
    ['Manager 主备交互支持', '补充交互引擎', 'v0.2'],
    ['sudo 脚本更新步骤', '新增部署步骤', 'v0.2'],
    ['Web UI 集群创建引导', '新增交互模式', 'v0.3'],
    ['详细验证清单', '补充验证阶段', 'v0.2'],
    ['多版本适配（8.x/9.x）', '版本参数化', 'v1.0'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')

doc.add_heading('8.4  更新后的完整部署全景（9阶段）', level=2)
doc.add_paragraph('综合两份文档分析，完整部署流程扩展为 9 个阶段：')
p = doc.add_paragraph()
run = p.add_run('阶段0:LLD规划 -> 阶段1:环境准备 -> 阶段2:OS调优 -> 阶段3:硬件验证 -> 阶段4:磁盘规划 -> 阶段5:网络验证 -> 阶段6:软件部署 -> 阶段7:预检查+Manager安装 -> 阶段8:Web集群部署+验证')
run.bold = True

doc.add_paragraph('')
doc.add_paragraph('交互助手应设计为版本参数化，根据检测到的 DWS 版本和 OS 类型自动适配步骤和参数。')

doc.add_paragraph('')

# ════════════════════════════════════════════
# 第九章  网上资源综合分析与建议
# ════════════════════════════════════════════
doc.add_heading('第九章  网上资源综合分析与建议', level=1)

doc.add_paragraph(
    '通过搜索华为云社区、CSDN、华为官方文档等渠道获取的 DWS 部署相关资源，'
    '与本文方案进行综合对比，识别出新的补充项和开发建议。'
)

# 9.1
doc.add_heading('9.1  核心发现：90% 安装问题是环境问题', level=2)
doc.add_paragraph(
    '根据华为云社区的官方问题案例统计，90% 的 DWS 安装问题属于环境问题。'
    '核心建议包括：安装前重装 OS、使用原生 ISO 镜像、不升级 RPM 包、提前检查网络和网卡配置。'
    '这进一步验证了前置全量检查（P0）的重要性。'
)

# 9.2 常见问题汇总
doc.add_heading('9.2  常见问题结构化汇总', level=2)
doc.add_paragraph('从华为云官方问题案例集及其他资源中提取的典型问题，按阶段分类：')

doc.add_heading('1. Preinstall 阶段', level=3)
headers = ['问题', '原因', '解决方案']
data = [
    ['RPM 补充安装失败', '使用非原生 ISO / 升级了 RPM 包', '使用原生 ISO，联系 OS 人员'],
    ['autopart 自动分区失败', '上传文件导致分区目录丢失', '重新上传配置，执行清理后重试'],
    ['挂盘失败', '磁盘 > 2TB 但文件系统为 ext3', '修改文件系统为 ext4 或 xfs'],
    ['分区失败', 'LLD 中元数据分区 RAID 类型不匹配', '修改为 noraid'],
]
add_styled_table(doc, headers, data)

doc.add_heading('2. OMS（Manager）安装阶段', level=3)
headers = ['问题', '原因', '解决方案']
data = [
    ['LDAP 版本不匹配', 'openldap RPM 包版本不对', '按产品文档安装对应版本'],
    ['nodeagent 启动失败', '老集群残留进程', '重装 OS 或执行卸载残留脚本'],
    ['备 OMS NTP 超时', 'LocalBackup 目录权限异常/文件残留', '清理目录后重试'],
    ['OMS 频繁主备倒换', '数据节点残留 nodeagent', '执行残留清理脚本'],
    ['证书不一致', '主备节点 cacert 不一致', '从主节点拷贝证书到备节点'],
]
add_styled_table(doc, headers, data)

doc.add_heading('3. 创建集群阶段', level=3)
headers = ['问题', '原因', '解决方案']
data = [
    ['校验请求参数失败', 'nodeagent 异常/hosts 权限异常', 'hosts 改为 644 权限'],
    ['初始化系统环境失败', 'omm 用户残留', 'userdel -rf omm，清理 uid_list'],
    ['分发软件包失败', 'SSH 互信异常/openssh 版本过高', '重建互信或回退 openssh'],
    ['安装节点失败', '磁盘空间不足', '清理日志目录释放空间'],
    ['初始化集群失败', '数据盘 > 20T（8.1.3 限制）', 'RAID 降至 20T 以下'],
    ['gs_install 失败', 'Python3 版本 < 3.7', '安装 Python 3.8.5'],
    ['OS 识别异常', '/etc/os-release 缺少引号', '手动添加引号'],
]
add_styled_table(doc, headers, data)

doc.add_heading('4. 其他常见问题', level=3)
doc.add_paragraph('Python 依赖库缺失：ldd /usr/bin/python3 检查缺少 libpython3.7m.so.1.0 等')
doc.add_paragraph('根目录磁盘使用率超过 90%：清理 /var/log 等目录')
doc.add_paragraph('Python 包权限异常：chmod -R 755 /usr/local/lib/python3.7/site-packages')
doc.add_paragraph('SSH 连接失败：检查 PermitRootLogin yes、PasswordAuthentication yes')

# 9.3
doc.add_heading('9.3  新增功能建议', level=2)

doc.add_heading('建议1：内置故障知识库', level=3)
doc.add_paragraph(
    '将上述常见问题结构化，嵌入交互助手中。当某个步骤执行失败时，'
    '自动匹配已知问题模式并显示可能的原因和解决方案，而不是简单地标记"失败"。'
)
doc.add_paragraph('实现方式：')
items = [
    '每个步骤关联一个故障列表（error_code / 输出模式匹配）',
    '失败时自动检索并显示 TOP 3 可能原因',
    '提供"查看华为云官方案例"的链接引导',
]
for item in items:
    doc.add_paragraph(f'  * {item}')

doc.add_heading('建议2：集成 gs_check 工具', level=3)
doc.add_paragraph(
    '华为提供了 gs_check 环境检查工具，可用于部署前和扩容前的巡检。'
    '交互助手应集成此工具而非完全重新实现检查逻辑。'
)
items = [
    '部署前调用 gs_check -e preinstall 进行环境检查',
    '扩容前调用 gs_check -e expand 进行扩容检查',
    '自动解析检查报告，可视化展示通过/失败项',
    '将结果纳入统一审计日志',
]
for item in items:
    doc.add_paragraph(f'  * {item}')

doc.add_heading('建议3：残留清理自动化', level=3)
doc.add_paragraph(
    'OMS 和集群安装失败的 #1 原因是老集群残留。交互助手应增加专门的残留检测和清理模块：'
)
items = [
    '检测 omm 用户是否存在（需清理）',
    '检测 nodeagent 进程是否运行',
    '检测 /etc/uid_list 是否存在',
    '检测 /srv/BigData 目录是否有残留',
    '检测 /opt/FusionInsight_* 目录是否有残留',
    '一键执行清理脚本',
]
for item in items:
    doc.add_paragraph(f'  * {item}')

doc.add_heading('建议4：症状诊断模式', level=3)
doc.add_paragraph(
    '当部署失败时，工具应支持"诊断模式"，用户选择症状后自动给出排查建议：'
)
headers = ['症状', '可能原因', '排查命令']
data = [
    ['nodeagent 启动失败', '老集群残留 / 权限异常', 'ps -ef|grep nodeagent; ls -la /opt/huawei'],
    ['备 OMS 安装超时', 'LocalBackup 目录异常', 'df -h /srv/BigData/LocalBackup; ls -la'],
    ['创建集群校验失败', 'hosts 权限 / omm 残留', 'ls -la /etc/hosts; id omm'],
    ['软件包分发失败', 'SSH 互信异常', 'ssh 节点IP "date" 测试'],
    ['gs_install 失败', 'Python 版本/依赖', 'python3 --version; ldd /usr/bin/python3'],
]
add_styled_table(doc, headers, data)

doc.add_heading('建议5：Python 环境深度检测', level=3)
doc.add_paragraph(
    '目前只检查 Python 版本号，但实际部署中常见问题是 Python 依赖库缺失。'
    '应增加深度检测：'
)
items = [
    'python3 --version 检查版本号',
    'ldd /usr/bin/python3 检查动态库依赖',
    '检查 /usr/local/lib/python3.*/site-packages 权限为 755',
    '验证 pip3 可用性（离线环境验证）',
    '检查 import sqlite3 / ctypes 等关键模块是否可用',
]
for item in items:
    doc.add_paragraph(f'  * {item}')

# 9.4
doc.add_heading('9.4  开发计划更新建议', level=2)
doc.add_paragraph('基于综合资源分析，建议调整开发计划如下：')

headers = ['版本', '新增/调整内容', '来源']
data = [
    ['v0.1', '前置全量检查 + Python 深度检测', '问题案例集'],
    ['v0.1', 'LLD 工具集成配置文件生成', '博客/文档'],
    ['v0.2', '内置故障知识库（分阶段归类）', '华为云问题案例集'],
    ['v0.2', '残留清理模块', '问题案例集'],
    ['v0.2', 'gs_check 工具集成', '华为官方文档'],
    ['v0.3', '症状诊断模式', '综合'],
    ['v0.3', '部署报告对接 ManageOne 验收', '华为官方文档'],
    ['v1.0', 'ManageOne 自动化验收集成', '华为官方文档'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')

# 9.5
doc.add_heading('9.5  推荐参考资源清单', level=2)

headers = ['资源名称', '来源', '说明']
data = [
    ['GaussDB(DWS)线下纯软/ESL集群安装问题案例集', '华为云社区', '最全的问题及解决方案汇总'],
    ['GaussDB分布式集群部署实战', 'CSDN', '从零搭建高可用环境'],
    ['GaussDB(DWS) 产品文档', '华为官方', '部署参数详细说明'],
    ['GaussDB(DWS) 配置规划工具', 'LLD Excel', '生成部署配置文件'],
    ['gs_check 巡检工具', '华为官方', '部署前/扩容前自动巡检'],
    ['ManageOne 自动化验收', '华为官方', '部署后一键验收+报告导出'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')

doc.add_page_break()

# ════════════════════════════════════════════
# 第十章  DWS 运维管理补充
# ════════════════════════════════════════════
doc.add_heading('第十章  DWS 运维管理补充', level=1)

doc.add_paragraph(
    '通过华为云社区、华为官方文档等渠道获取的 DWS 运维管理相关资料，'
    '补充以下运维场景的功能建议。'
)

# 10.1
doc.add_heading('10.1  集群升级管理', level=2)

doc.add_paragraph('DWS ESL 版本升级流程及时间规划：')

headers = ['阶段', '内容', '参考时长']
data = [
    ['升级前巡检', 'Tool Prober 检查前置条件', '—'],
    ['升级前准备', '处理告警、部署升级工具', '—'],
    ['升级集群', 'UpdateService 自动化升级（需停业务）', '约 2 小时'],
    ['升级后操作', '检查、清理升级工具、升级客户端驱动', '约 2 小时'],
    ['建议总窗口', '含 4 小时应急预案', '约 8 小时预留'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')
doc.add_paragraph('升级方式演进：大版本全量升级 → 就地升级 → 小版本离线 → 小版本滚动 → 闪断升级（当前，仅闪断一次业务）')

doc.add_paragraph('常见升级问题：')
upgrade_issues = [
    'tmp 目录 t 权限导致解密失败（麒麟 V10 等欧拉系）',
    'openssl 版本不一致导致证书安装失败',
    'SSH 互信异常（密钥加密字段缺失）',
    '目录软链接导致路径不一致',
    'ipv6 未开启导致端口绑定失败',
    '系统表依赖导致升级报错（用户视图依赖系统视图）',
]
for item in upgrade_issues:
    doc.add_paragraph(f'  * {item}')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('对交互助手的意义：增加升级管理模块，覆盖升级前巡检→升级引导→回滚支持→后验证全流程，集成升级问题诊断知识库。')
run.bold = True

# 10.2
doc.add_heading('10.2  巡检体系', level=2)

headers = ['工具', '用途', '适用场景']
data = [
    ['FusionInsight Tool Prober', '升级前全量前置检查', '升级/补丁前'],
    ['gs_check', '集群巡检命令，支持自定义场景', '日常运维、扩容前'],
    ['自定义巡检场景', 'scene_XXX.xml + items.xml 配置', '特定业务检查'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')
doc.add_paragraph('巡检项分类：')
headers = ['分类', '巡检项示例']
data = [
    ['OS', 'CPU 占用率、内存使用率、磁盘使用率、文件句柄数'],
    ['Network', '网络丢包率、端口范围、网卡带宽'],
    ['Cluster', '集群拓扑、节点一致性、关键进程状态'],
    ['Database', '系统表膨胀、表倾斜、锁数量、活跃连接数、备份结果、归档参数、慢 SQL、Vacuum 状态'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('对交互助手的意义：集成 gs_check 巡检工具，自动触发巡检并解析结果，替代手动执行和人工阅读日志。')
run.bold = True

# 10.3
doc.add_heading('10.3  告警管理', level=2)
doc.add_paragraph('DWS ESL 版本默认包含 19 个告警项，归类如下：')

headers = ['告警类别', '告警项示例', '阈值']
data = [
    ['磁盘使用率', '系统盘/日志盘/数据盘使用率超阈值', '紧急 >85%，重要 >80%'],
    ['CPU 使用率', '节点 CPU 使用率超阈值', '系统+用户 CPU'],
    ['I/O 异常', 'I/O 利用率/时延超阈值', '紧急'],
    ['inode 耗尽', '系统盘/日志盘/数据盘 inode 使用率', '紧急 >95%，重要 >90%'],
    ['查询异常', '下盘量超阈值/查询堆积超阈值', '紧急'],
    ['资源池阻塞', '默认资源池队列阻塞', '紧急'],
    ['SQL 探针', '探针耗时超阈值', '8.1.1.300+ 版本'],
    ['长锁 Vacuum', '持有表锁过长的 Vacuum Full', '重要，8.2.0.100+'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')
doc.add_paragraph('告警订阅方式：支持短信、电子邮件、应用推送。推荐重点订阅磁盘使用率、CPU 使用率、查询下盘量、SQL 探针耗时、资源池阻塞。')

p = doc.add_paragraph()
run = p.add_run('对交互助手的意义：增加告警解读模块，部署后可自动配置推荐告警阈值和订阅规则。')
run.bold = True

# 10.4
doc.add_heading('10.4  扩容/缩容运维', level=2)

headers = ['操作', '版本要求', '业务影响', '注意事项']
data = [
    ['在线扩容', '8.1.1+', '不中断，重分布影响性能', '提前建快照、执行 VACUUM、清理临时表'],
    ['在线缩容', '8.1.1.300+', '不中断，数据搬迁', 'node 编号可能不连续，需刷新'],
    ['节点故障恢复', '—', '集群可能变只读', '删除复制槽、转移 xlog、重建备机'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')
doc.add_paragraph('扩容前最佳实践：')
items = [
    '创建手动快照备份数据',
    '执行 VACUUM 清理脏数据，释放存储空间',
    '关闭创建临时表的客户端连接',
    '扩容期间停止业务或仅运行少量查询',
    '表重分布期间避免执行超过 20 分钟的查询（写锁默认超时时间）',
    '大规模集群建议关闭自动重分布，完成后手动执行以支持失败重试',
]
for item in items:
    doc.add_paragraph(f'  * {item}')

doc.add_paragraph('')
doc.add_paragraph('典型案例：缩容后快照失败 — node 编号不连续导致备份数组越界，需用 Python 脚本刷新 node 编号为连续 ID。')

p = doc.add_paragraph()
run = p.add_run('对交互助手的意义：增加扩缩容引导模块，自动执行前置检查（快照/VACUUM/状态确认）、引导执行、后置验证（node 编号/快照可用性）。')
run.bold = True

# 10.5
doc.add_heading('10.5  备份恢复', level=2)
doc.add_paragraph('DWS ESL 的备份恢复机制：')
items = [
    '全量备份 + 增量备份：巡检会检查近两天的备份是否正常',
    '升级时系统表自动备份：global, pg_clog, pg_xlog, pg_multixact, pg_replslot, pg_notify, pg_subtrans, pg_twophase',
    '主备切换：DN 主故障后自动切换备 DN 升主，RPO=0, RTO<10s（同城 AZ 内）',
    '容灾方案：同城跨 AZ 双活、两地三中心异地容灾、跨 Region 数据实时灾备',
]
for item in items:
    doc.add_paragraph(f'  * {item}')

p = doc.add_paragraph()
run = p.add_run('对交互助手的意义：增加备份检查模块，部署后自动验证备份策略是否生效，升级前自动触发系统表备份。')
run.bold = True

# 10.6
doc.add_heading('10.6  运维阶段日志路径知识库', level=2)
doc.add_paragraph('各运维场景的关键日志路径，用于诊断模块自动关联：')

headers = ['运维场景', '日志路径', '节点位置']
data = [
    ['升级/补丁', '/var/log/bigdata/update-service', '主 OMS 节点'],
    ['Manager 运行', '/var/log/bigdata/controller', '主 OMS 节点'],
    ['数据库升级', '/var/log/bigdata/mpp/upgrade', '主 CMS 节点'],
    ['内核操作', '/var/log/Bigdata/mpp/omm/om/gs_upgradectl-*.log', '各节点'],
    ['节点本地', '/var/log/Bigdata/mpp/omm/om/gs_local-*.log', '故障节点'],
    ['安装部署', '/var/log/Bigdata/controller/controller.log', '主 OMS 节点'],
    ['MPPDB 安装', '/var/log/Bigdata/mpp/scriptlog/postinstall.log', '故障节点'],
    ['NodeAgent', '/var/log/Bigdata/nodeagent', '各 Agent 节点'],
]
add_styled_table(doc, headers, data)

# 10.7
doc.add_heading('10.7  新增功能建议汇总', level=2)

headers = ['运维场景', '建议功能', '优先级', '说明']
data = [
    ['升级管理', '升级全流程引导 + 回滚支持', 'P1', '覆盖巡检→升级→验证→回滚'],
    ['巡检管理', '集成 gs_check + 结果可视化', 'P1', '自动触发巡检并解析报告'],
    ['告警管理', '告警阈值推荐 + 订阅配置', 'P2', '部署后自动配置推荐规则'],
    ['扩容/缩容', '前置检查 + 执行引导 + 后置验证', 'P1', '快照/VACUUM/node编号检查'],
    ['备份恢复', '备份状态检查 + 升级前自动备份', 'P2', '验证备份策略是否生效'],
    ['日志诊断', '故障时自动关联对应日志路径', 'P1', '按场景提示查看对应日志文件'],
]
add_styled_table(doc, headers, data)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('以上运维功能建议纳入 v0.3 之后的开发计划，使交互助手从"部署工具"扩展为"部署+运维一体化助手"。')
run.bold = True

doc.add_paragraph('')

# 结尾
p = doc.add_paragraph()
run = p.add_run('附：原始文档基于 ')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = p.add_run('DWS部署-V2.0.pdf')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = p.add_run('（华为 FusionInsight DWS 9.1.0 on KylinOS ARM64）')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run = p.add_run(' 及博客《麒麟V10 SP2安装GaussDB(DWS)-ESL》')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('参考资源：')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.bold = True
resources = [
    '华为云官方问题案例集 — GaussDB(DWS)线下纯软/ESL集群安装问题案例集',
    'Huawei FusionInsight DWS 产品文档',
    'GaussDB(DWS) 8.2.1-ESL 配置规划工具（LLD Excel）',
]
for r in resources:
    p = doc.add_paragraph()
    run = p.add_run(f'  * {r}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ════════════════════════════════════════════
# 保存
# ════════════════════════════════════════════
output = os.path.expanduser('~/Desktop/DWS部署交互助手-开发建议书_v3.docx')
doc.save(output)
print(f'Saved: {output}')

# 文件大小
size = os.path.getsize(output)
print(f'Size: {size / 1024:.0f} KB')
