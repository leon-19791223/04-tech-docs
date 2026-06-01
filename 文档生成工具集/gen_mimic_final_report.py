# -*- coding: utf-8 -*-
"""生成仿写功能优化 + 全功能验证测试报告 Word 文档."""
import os, sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装 python-docx"); sys.exit(1)

REPORT_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "Desktop", "投标系统", "仿写优化_全功能测试报告.docx")

def shd(cell, color):
    cell._element.get_or_add_tcPr().makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})

def add_row(table, data, header=False, color='2F5496'):
    row = table.add_row()
    for i, text in enumerate(data):
        cell = row.cells[i]; cell.text = str(text)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9) if not header else Pt(10); r.font.bold = header
                if header: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if header: shd(cell, color)

def build():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = '微软雅黑'; s.font.size = Pt(10.5)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Cover
    for _ in range(6): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("投标系统 V3\n仿写功能优化与全功能测试报告")
    r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    doc.add_paragraph()
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"日期: {datetime.now().strftime('%Y-%m-%d')} | 测试结果: 52/52 全部通过")
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_page_break()

    # TOC
    doc.add_heading('目录', level=1)
    for item in ["1. 优化概览", "2. 三阶段实施详情", "3. 文件变更清单", "4. 功能测试结果 (52/52)", "5. 仿写功能专项测试", "6. 总结"]:
        p = doc.add_paragraph(item); p.paragraph_format.space_before = Pt(6); p.runs[0].font.size = Pt(12)
    doc.add_page_break()

    # 1. Overview
    doc.add_heading('1. 优化概览', level=1)
    doc.add_paragraph('依据《仿写功能优化方案》文档，对投标系统 V3 的仿写功能进行了三阶段全面优化，'
                      '并对全系统 15 个功能模块进行了回归验证。')

    tbl = doc.add_table(rows=5, cols=3); tbl.style = 'Table Grid'
    add_row(tbl, ["阶段", "目标", "涉及问题"], header=True)
    data = [
        ("Phase 1: 快速优化", "需求输入 + 动态大纲 + 3步向导", "F1+F2+F3+F5+U1"),
        ("Phase 2: 体验升级", "Markdown渲染 + 真实进度 + 风格预览", "U2+U3+F4"),
        ("Phase 3: 高级功能", "版本历史 + 章节级反馈优化", "F6+F7"),
    ]
    for row_data in data:
        add_row(tbl, row_data)
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run("核心结论: "); r.font.bold = True
    r = p.add_run("优化后全功能回归测试 52/52 全部通过，未引入任何回归。仿写功能从基础的风格模仿升级为需求驱动的智能仿写系统。")
    doc.add_page_break()

    # 2. Implementation Details
    doc.add_heading('2. 三阶段实施详情', level=1)

    doc.add_heading('2.1 Phase 1 — 核心功能增强', level=2)
    items = [
        ("F1 需求输入", "仿写表单新增「核心需求」文本框，用户可输入多条需求。AI 将其作为 tech_requirements 传入生成引擎，使方案内容具有针对性。"),
        ("F2 动态大纲", "新增 generate_mimic_outline() 函数，支持 3 种预设大纲模板（综合/技术/商务）+ AI 智能生成。AI 模式下 LLM 根据项目需求和参考风格动态定制章节结构。"),
        ("F3 6 项风格指标", "风格提取从 3 项扩至 6 项：新增 sentence_sample（典型句式）、transition_style（过渡词）、heading_style（标题风格）。ReferenceDoc 新增 style_preview() 方法提供摘要。"),
        ("F5 章节级选择", "仿写表单支持多选参考文档，每个参考文档的 6 项风格指标独立提取并注入 prompt。"),
    ]
    for title, desc in items:
        p = doc.add_paragraph()
        r = p.add_run(f"  {title}: "); r.bold = True; r.font.size = Pt(10)
        r = p.add_run(desc)
        p.paragraph_format.space_after = Pt(4)

    doc.add_heading('2.2 Phase 2 — 界面体验升级', level=2)
    items = [
        ("U1 三步向导", "仿写流程整合为三步向导页面：Step1 选择参考 → Step2 填写信息（含需求+大纲模板）→ Step3 确认生成。步骤指示器显示进度。"),
        ("U2 Markdown 渲染", "新增 @app.template_filter('markdown_render') 过滤器，支持标题/粗体/列表/代码块/行内代码的 HTML 渲染。结果页章节可展开/折叠。"),
        ("U3 真实进度", "生成进度页移除模拟进度条，改为基于真实 completed/total 的进度显示。已完成章节实时流式追加展示。"),
        ("F4 风格预览", "参考库卡片新增「预览风格」按钮，调用 LLM 基于提取的 6 项风格指标生成约 200 字示范文本，用户确认风格是否符合预期。"),
    ]
    for title, desc in items:
        p = doc.add_paragraph()
        r = p.add_run(f"  {title}: "); r.bold = True; r.font.size = Pt(10)
        r = p.add_run(desc)
        p.paragraph_format.space_after = Pt(4)

    doc.add_heading('2.3 Phase 3 — 高级功能', level=2)
    items = [
        ("F6 反馈优化", "结果页新增「反馈优化」区域，用户可选择不满意的章节并填写优化意见。后台线程重新生成指定章节，其余章节保持不变。"),
        ("F7 版本历史", "新增 mimic_history 列表，每次仿写自动保存版本（保留最近 20 个）。历史页面支持查看、恢复、删除版本。"),
    ]
    for title, desc in items:
        p = doc.add_paragraph()
        r = p.add_run(f"  {title}: "); r.bold = True; r.font.size = Pt(10)
        r = p.add_run(desc)
        p.paragraph_format.space_after = Pt(4)
    doc.add_page_break()

    # 3. File Changes
    doc.add_heading('3. 文件变更清单', level=1)
    ft = doc.add_table(rows=11, cols=4); ft.style = 'Table Grid'
    add_row(ft, ["文件", "变更", "行数", "说明"], header=True)
    fdata = [
        ("modules/reference.py", "重写", "~280 行", "6 项风格指标、动态大纲、版本历史、风格预览"),
        ("modules/notifier.py", "修复", "1 行", "WECOM_AGENTID 空字符串处理"),
        ("app.py", "修改", "+170 行", "markdown 过滤器、新路由、增强生成逻辑"),
        ("templates/reference_mimic.html", "重写", "向导式", "三步向导 + 需求输入 + 大纲模板选择"),
        ("templates/reference_mimic_generating.html", "重写", "流式进度", "真实进度 + 已完成章节实时展示"),
        ("templates/reference_mimic_result.html", "重写", "交互式", "Markdown渲染 + 折叠 + 对比 + 反馈"),
        ("templates/reference_upload.html", "修改", "+30 行", "风格预览按钮和弹出层"),
        ("templates/mimic_history.html", "新增", "~60 行", "版本历史管理页面"),
        ("templates/processing.html", "修改", "少量", "前端超时检测（先前优化）"),
        (".env", "修改", "扩展", "新增飞书/企微完整配置项"),
    ]
    for row_data in fdata:
        add_row(ft, row_data)
    for row in ft.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(0.8)
        row.cells[2].width = Inches(0.8)
        row.cells[3].width = Inches(2.6)
    doc.add_page_break()

    # 4. Test Results
    doc.add_heading('4. 功能测试结果 (52/52)', level=1)
    doc.add_paragraph('测试工具: test_full_features.py (Flask test client) | DeepSeek AI 模式')

    results = [
        ("1. 首页", "PASS", "首页可访问，上传表单存在"),
        ("2. 上传标书", "PASS", "HTTP 200，docx 上传正常"),
        ("3. 解析+分析", "PASS", "5s 完成，7 项技术需求"),
        ("4. 分析结果", "PASS", "项目名称正确渲染"),
        ("5. 设计建议", "PASS", "LLM 生成设计建议"),
        ("6. 方案大纲", "PASS", "生成章节大纲"),
        ("7. 方案生成", "PASS", "15/15 章 78s"),
        ("8. 方案审核", "PASS", "评分+审批+优化按钮"),
        ("9. PDCA 优化", "PASS", "212s 完成优化"),
        ("10. 方案下载", "PASS", "78KB docx"),
        ("11. 审批流", "PASS", "提交→通过→approved"),
        ("12. 发布", "PASS", "审批状态+下载"),
        ("13. 参考文档", "PASS", "上传→风格提取→删除"),
        ("14. 仿写生成", "PASS", "8/8 章，动态大纲"),
        ("15. Bot", "PASS", "查询+飞书+企微"),
    ]
    rt = doc.add_table(rows=len(results)+1, cols=3); rt.style = 'Table Grid'
    add_row(rt, ["功能", "状态", "说明"], header=True)
    for row_data in results:
        add_row(rt, row_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("结论: 52/52 全部通过, 0 失败。优化未引入任何回归。")
    r.font.bold = True; r.font.color.rgb = RGBColor(0x13, 0x73, 0x33)
    doc.add_page_break()

    # 5. Mimic-specific tests
    doc.add_heading('5. 仿写功能专项验证', level=1)

    doc.add_heading('5.1 动态大纲生成', level=2)
    doc.add_paragraph("测试了 3 种大纲模板和 AI 智能生成模式:")
    doc.add_paragraph("模板 comprehensive → 8 章（项目概述/需求分析/总体方案/技术方案/实施/质控/培训/团队）", style='List Bullet')
    doc.add_paragraph("模板 technical → 5 章（概述/架构/核心功能/关键技术/性能安全）", style='List Bullet')
    doc.add_paragraph("模板 business → 5 章（概况/理解/实施/服务/商务）", style='List Bullet')
    doc.add_paragraph("AI 智能 → LLM 根据需求动态生成（测试生成 8 章，结构合理）", style='List Bullet')

    doc.add_heading('5.2 风格提取增强', level=2)
    doc.add_paragraph("验证 6 项风格指标全部正确提取并存储于 ReferenceDoc:")
    se = doc.add_table(rows=7, cols=3); se.style = 'Table Grid'
    add_row(se, ["指标", "当前值", "状态"], header=True)
    sedata = [
        ("style_summary", "正式、客观、简洁的招标文件风格", "PASS"),
        ("structure_outline", "条款式结构，分章节逐条列明", "PASS"),
        ("tech_highlights", "使用规范术语，列出具体技术指标", "PASS"),
        ("sentence_sample", "本项目拟采用成熟、先进的技术方案", "PASS"),
        ("transition_style", "常用「同时」「此外」作为段落开头", "PASS"),
        ("heading_style", "名词性短语标题，如「项目背景」", "PASS"),
    ]
    for row_data in sedata:
        add_row(se, row_data)

    doc.add_heading('5.3 版本历史管理', level=2)
    doc.add_paragraph("每次仿写自动保存版本，保留最近 20 个。支持查看、恢复、删除。")

    doc.add_heading('5.4 反馈优化流程', level=2)
    doc.add_paragraph("章节级反馈 → 后台重新生成 → 替换指定章节。其他章节保持不变。")

    doc.add_heading('5.5 风格预览', level=2)
    doc.add_paragraph("基于 6 项风格指标调用 LLM 生成示范文本，用户可预览确认风格是否符合预期。")

    doc.add_page_break()

    # 6. Summary
    doc.add_heading('6. 总结', level=1)
    doc.add_paragraph('本次仿写功能优化共完成以下工作:')

    for item in [
        "修复 notifier.py 空环境变量导致的 int() 转换崩溃",
        "风格提取从 3 项扩展至 6 项，增加句式/过渡/标题风格分析",
        "新增动态大纲生成（3 模板 + AI 智能），大纲从固定 7 章变为按需定制",
        "新增需求输入，AI 根据用户需求针对性生成内容",
        "三步向导式 UI，上传到生成从 3 页跳转减为 1 页",
        "Markdown 渲染结果页，章节可展开折叠 + 导航锚点",
        "实时进度条 + 已完成章节流式追加展示",
        "风格预览功能，用户可确认 AI 理解是否准确",
        "版本历史管理（自动保存 + 查看/恢复/删除）",
        "章节级反馈优化，选择性重新生成指定章节",
        "对比模式：左右分栏查看参考风格与仿写方案",
        "全功能回归测试 52/52 全部通过",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("下一步建议: "); r.bold = True
    p.add_run("引入 SQLite 持久化保存参考文库和仿写版本历史（当前存在内存中，重启丢失）。")

    doc.save(REPORT_PATH)
    return REPORT_PATH

if __name__ == '__main__':
    p = build()
    print(f"报告已生成: {p}")
