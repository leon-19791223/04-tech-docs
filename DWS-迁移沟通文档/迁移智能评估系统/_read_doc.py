import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
path = os.path.expanduser(r'~\Desktop\Z-DBMate 用户手册（for KylinV10） V1.0.docx')
doc = Document(path)
print(f'段落: {len(doc.paragraphs)}, 表格: {len(doc.tables)}')
print('='*60)

level_map = {
    'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3,
    'Heading 4': 4, 'Heading 5': 5, 'Heading 6': 6,
}

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    s = p.style.name if p.style else 'Normal'
    if s in level_map:
        lv = level_map[s]
        prefix = '#' * lv
        print(f'{prefix} {t}')
    else:
        print(t[:200])
