import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
path = os.path.expanduser(r'~\Desktop\Z-DBMate 用户手册（for KylinV10） V1.0.docx')
doc = Document(path)

print(f'表格数量: {len(doc.tables)}')
print('='*60)

for ti, table in enumerate(doc.tables):
    print(f'\n--- 表格 {ti+1} ({len(table.rows)}行 x {len(table.columns)}列) ---')
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n',' ') for cell in row.cells]
        print(' | '.join(cells))
