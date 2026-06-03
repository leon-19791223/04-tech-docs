"""
GP → DWS 迁移智能评估报告生成器
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from core.models import AssessmentResult

# ============================================================
# 颜色常量
# ============================================================
BLUE_DARK = RGBColor(0x1A, 0x3C, 0x6E)
BLUE = RGBColor(0x00, 0x7A, 0xCC)
DARK = RGBColor(0x2D, 0x2D, 0x2D)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
ORANGE = RGBColor(0xF5, 0x7F, 0x17)
RED = RGBColor(0xC6, 0x28, 0x28)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_header_row(table, headers, bg_color="1A3C6E"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, bg_color)

def add_data_row(table, values, bold_first=True):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(v))
        run.font.size = Pt(9)
        run.font.color.rgb = DARK
        if i == 0 and bold_first:
            run.bold = True
    return row

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE_DARK
    return h

def add_para(doc, text, bold=False, color=None, size=10, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or DARK
    run.bold = bold
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_rich_para(doc, segments, space_after=6):
    p = doc.add_paragraph()
    for seg in segments:
        text, bold, color = seg[0], seg[1] if len(seg) > 1 else False, seg[2] if len(seg) > 2 else DARK
        size = seg[3] if len(seg) > 3 else 10
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_score_box(doc, score, risk_level):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{score}")
    run.font.size = Pt(48)
    run.bold = True
    run.font.color.rgb = GREEN if score >= 85 else ORANGE if score >= 65 else RED
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("兼容性评分")
    run2.font.size = Pt(12)
    run2.font.color.rgb = GRAY
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    level_map = {"低": "低风险", "中": "中风险", "高": "高风险"}
    run3 = p3.add_run(f"风险等级: {level_map.get(risk_level, risk_level)}")
    run3.font.size = Pt(14)
    run3.bold = True
    run3.font.color.rgb = GREEN if risk_level == "低" else ORANGE if risk_level == "中" else RED
    bg = "E8F5E9" if score >= 85 else "FFF8E1" if score >= 65 else "FFEBEE"
    set_cell_shading(cell, bg)
    return table


class ReportGenerator:
    """评估报告生成器"""

    def __init__(self, result: AssessmentResult, output_path: str):
        self.result = result
        self.output_path = output_path

    def generate(self) -> str:
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        self._add_cover(doc)
        self._add_summary(doc)
        self._add_environment_overview(doc)
        self._add_compatibility_detail(doc)
        self._add_critical_issues(doc)
        self._add_workload_estimate(doc)
        self._add_phase_breakdown(doc)
        self._add_phase_timeline(doc)
        self._add_capacity_planning(doc)
        self._add_batch_strategy(doc)
        self._add_recommendations(doc)
        self._add_appendix(doc)
        doc.save(self.output_path)
        return self.output_path

    def _add_cover(self, doc):
        r = self.result
        path_label = f"{r.source_type.upper() if r.source_type else 'GP'} → GaussDB(DWS)"
        for _ in range(6):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(path_label)
        run.font.size = Pt(28)
        run.font.color.rgb = BLUE_DARK
        run.bold = True
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("迁移智能评估报告")
        run2.font.size = Pt(22)
        run2.font.color.rgb = BLUE
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run("━" * 30)
        run3.font.color.rgb = BLUE
        run3.font.size = Pt(12)
        doc.add_paragraph()
        doc.add_paragraph()
        for item in [
            f"源端数据库: {r.db_version or '—'}",
            f"目标数据库: GaussDB(DWS) 8.2.1",
            f"集群规模: {r.cluster_scale or '—'}",
            f"数据总量: {r.total_capacity or '—'}",
            f"源库字符集: {self._get_meta('source_charset', '—')}",
            f"评估日期: {datetime.now().strftime('%Y-%m-%d')}",
            f"文档版本: v1.0",
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item)
            run.font.size = Pt(12)
            run.font.color.rgb = GRAY
        doc.add_page_break()

    def _get_meta(self, key, default=""):
        """获取元数据属性"""
        from core.models import MigrationMetadata
        return getattr(self.result, key, default)

    def _add_summary(self, doc):
        add_heading(doc, "1  评估概要", level=1)
        doc.add_paragraph()
        add_score_box(doc, self.result.overall_score, self.result.risk_level)
        doc.add_paragraph()
        add_heading(doc, "1.1  关键指标", level=2)
        r = self.result
        cat_scores = {c.category: c for c in r.category_results}
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        make_header_row(table, ["评估维度", "得分", "状态", "关键项"])
        summary_map = {
            "ddl": "DDL兼容性",
            "data_type": "数据类型兼容性",
            "function": "函数兼容性",
            "udf_language": "UDF语言兼容性",
            "plsql": "PL/SQL语法兼容性",
            "sqlpl": "SQL PL存储过程兼容性",
            "trigger": "触发器兼容性",
            "extension": "扩展兼容性",
            "exadata": "Oracle Exadata特性兼容性",
            "etl_tool": "ETL工具兼容性",
            "scheduler": "调度工具兼容性",
            "bi_tool": "BI工具兼容性",
            "security": "安全与权限兼容性",
            "charset": "字符集与编码兼容性",
            "app_layer": "应用层兼容性",
            "transaction": "事务与并发兼容性",
            "cdc": "CDC增量同步兼容性",
            "performance": "性能与容量兼容性",
        }
        for cat_id, cat_name in summary_map.items():
            cat = cat_scores.get(cat_id)
            if cat is None:
                continue
            row = add_data_row(table, [
                cat_name,
                f"{cat.actual_score}",
                "✅ 兼容" if cat.errors == 0 else f"⚠️ {cat.errors}项不兼容",
                f"{cat.warnings}项需注意" if cat.warnings > 0 else "无",
            ])
            if cat.errors > 0:
                set_cell_shading(row.cells[2], "FFEBEE")
            elif cat.warnings > 0:
                set_cell_shading(row.cells[2], "FFF8E1")
            else:
                set_cell_shading(row.cells[2], "E8F5E9")
        doc.add_paragraph()
        add_heading(doc, "1.2  总体结论", level=2)
        score = self.result.overall_score
        src_type = self.result.source_type.upper() if self.result.source_type else "源端"
        if score >= 85:
            conclusion = (
                f"评估综合得分 {score} 分，迁移风险等级为【低风险】。"
                f"{src_type}与DWS整体兼容性良好，主要差异集中在DDL语法微调和少数不兼容功能。"
                f"建议按标准迁移流程分阶段推进。"
            )
        elif score >= 65:
            conclusion = (
                f"评估综合得分 {score} 分，迁移风险等级为【中风险】。"
                f"存在{len(self.result.critical_issues)}项不兼容问题和若干需关注项，"
                f"主要风险集中在存储过程/UDF语言兼容性和扩展替代。建议在POC阶段充分验证后推进。"
            )
        else:
            conclusion = (
                f"评估综合得分 {score} 分，迁移风险等级为【高风险】。"
                f"存在{len(self.result.critical_issues)}项严重不兼容问题，"
                f"建议进行专项技术评估，必要时调整迁移策略或分阶段实施。"
            )
        add_para(doc, conclusion, size=11)
        doc.add_page_break()

    def _add_environment_overview(self, doc):
        add_heading(doc, "2  环境总览", level=1)
        doc.add_paragraph()
        add_heading(doc, "2.1  源端环境信息", level=2)
        r = self.result
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        make_header_row(table, ["项目", "内容"])
        for k, v in [
            ("数据库版本", r.db_version or "—"),
            ("内核版本", r.kernel_version or "—"),
            ("集群规模", r.cluster_scale or "—"),
            ("总数据容量", r.total_capacity or "—"),
            ("ETL工具", r.etl_tool or "—"),
            ("调度工具", r.scheduler_tool or "—"),
            ("BI工具", ", ".join(r.bi_tools) if r.bi_tools else "—"),
        ]:
            add_data_row(table, [k, v])
        doc.add_paragraph()
        add_heading(doc, "2.2  元数据对象统计", level=2)
        table2 = doc.add_table(rows=1, cols=3)
        table2.style = 'Table Grid'
        make_header_row(table2, ["对象类型", "数量", "备注"])
        for k, v, note in [
            ("表", r.table_count, f"其中分区表{r.partition_table_count}张"),
            ("视图", r.view_count, ""),
            ("函数/UDF", r.function_count, "含多语言"),
        ]:
            add_data_row(table2, [k, str(v) if v else "—", note])
        doc.add_paragraph()
        if r.udf_languages:
            add_heading(doc, "2.3  UDF语言分布", level=2)
            table3 = doc.add_table(rows=1, cols=3)
            table3.style = 'Table Grid'
            make_header_row(table3, ["语言", "数量", "兼容性"])
            for lang, cnt in sorted(r.udf_languages.items(), key=lambda x: -x[1]):
                compat = "✅ 兼容" if lang in ("plpgsql", "sql") else "❌ 不兼容"
                add_data_row(table3, [lang, str(cnt), compat])
        doc.add_page_break()

    def _add_compatibility_detail(self, doc):
        add_heading(doc, "3  兼容性详细分析", level=1)
        doc.add_paragraph()
        order = 1
        for cat in self.result.category_results:
            if not cat.details:
                continue
            status_icon = "✅" if cat.errors == 0 else "⚠️"
            add_heading(doc, f"3.{order}  {status_icon} {cat.category_name}", level=2)
            add_para(doc, f"得分: {cat.actual_score}/100 | "
                          f"已检查{cat.total_rules}项 | "
                          f"通过{cat.passed}项 | "
                          f"警告{cat.warnings}项 | "
                          f"不兼容{cat.errors}项",
                     color=GRAY, size=9)
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            make_header_row(table, ["规则", "状态", "说明", "DWS解决方案"])
            for detail in cat.details:
                status_text = "✅ 兼容" if detail.severity == "info" else "⚠️ 需注意" if detail.severity == "warning" else "❌ 不兼容"
                row = add_data_row(table, [
                    detail.name,
                    status_text,
                    detail.description[:80] + ("..." if len(detail.description) > 80 else ""),
                    detail.target_solution[:60] + ("..." if len(detail.target_solution) > 60 else ""),
                ])
                if detail.severity == "error":
                    set_cell_shading(row.cells[1], "FFEBEE")
                elif detail.severity == "warning":
                    set_cell_shading(row.cells[1], "FFF8E1")
            doc.add_paragraph()
            order += 1
        doc.add_page_break()

    def _add_critical_issues(self, doc):
        add_heading(doc, "4  关键问题清单", level=1)
        doc.add_paragraph()
        if not self.result.critical_issues:
            add_para(doc, "✅ 未发现严重不兼容问题", color=GREEN, size=11, bold=True)
            return
        issues = self.result.critical_issues
        add_para(doc, f"共发现 {len(issues)} 项不兼容的关键问题，需在迁移前逐项解决:", color=RED, size=11, bold=True)
        doc.add_paragraph()
        for i, issue in enumerate(issues, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"问题 {i}: {issue['name']} [{issue['category']}]")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RED
            if issue.get('description'):
                add_para(doc, f"  描述: {issue['description']}", color=DARK, size=9)
            if issue.get('solution'):
                add_para(doc, f"  方案: {issue['solution']}", color=BLUE, size=9)
            if issue.get('difficulty'):
                add_para(doc, f"  改造难度: {issue['difficulty']}", color=GRAY, size=9)
            doc.add_paragraph()
        doc.add_page_break()

    def _add_workload_estimate(self, doc):
        add_heading(doc, "5  迁移工作量评估", level=1)
        doc.add_paragraph()
        wl = self.result.workload_estimate
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        make_header_row(table, ["工作项", "预估天数", "说明"])
        for name, days, note in [
            ("DDL转换", wl.get("ddl_conversion_days", "—"), "表结构/视图/索引DDL语法转换"),
            ("数据迁移", wl.get("data_migration_days", "—"), "全量+增量数据迁移及校验"),
            ("UDF改造", wl.get("udf_migration_days", "—"), "不兼容UDF语言改造与测试"),
            ("ETL迁移", wl.get("etl_migration_days", "—"), "ETL任务迁移适配"),
            ("测试验证", wl.get("testing_days", "—"), "功能测试+性能测试+回归测试"),
        ]:
            add_data_row(table, [name, str(days), note])
        doc.add_paragraph()
        add_rich_para(doc, [
            ("预估总工时: ", True, DARK, 12),
            (f"{wl.get('total_estimated_days', '—')} 人天 "
             f"(约 {wl.get('total_estimated_months', '—')} 人月)", True, BLUE, 12),
        ])
        if wl.get("note"):
            add_para(doc, f"注: {wl['note']}", color=GRAY, size=9)
        doc.add_paragraph()

        # 补充工作量明细(安全/字符集/应用层/事务/CDC)
        add_heading(doc, "5.1  补充评估工作量", level=2)
        table2 = doc.add_table(rows=1, cols=3)
        table2.style = 'Table Grid'
        make_header_row(table2, ["评估维度", "预估天数", "说明"])
        for name, days, note in [
            ("安全与权限改造", wl.get("security_audit_days", "—"), "权限迁移/LBAC改造/审计迁移"),
            ("字符集转换", wl.get("charset_conversion_days", "—"), "VARCHAR长度调整/编码转换/字符集测试"),
            ("应用层适配", wl.get("app_layer_adapt_days", "—"), "JDBC驱动替换/ORM方言变更/连接池调整"),
            ("事务改造", wl.get("transaction_adapt_days", "—"), "XA替代/自治事务改造/隔离级别验证"),
            ("CDC增量同步", wl.get("cdc_sync_days", "—"), "CDC工具配置/增量链路搭建/延迟监控"),
        ]:
            add_data_row(table2, [name, str(days), note])
        doc.add_page_break()

    def _add_phase_timeline(self, doc):
        add_heading(doc, "8  迁移阶段规划", level=1)
        doc.add_paragraph()
        phases = self.result.estimated_phases
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        make_header_row(table, ["阶段", "名称", "周期", "主要工作"])
        for key in sorted(phases.keys()):
            if key == "total_months":
                continue
            phase = phases[key]
            add_data_row(table, [
                f"Phase {phase.get('order', '—')}",
                phase.get("name", ""),
                f"{phase.get('weeks', '—')}周",
                phase.get("description", ""),
            ])
        doc.add_paragraph()
        if phases.get("total_months"):
            add_rich_para(doc, [
                ("预计总工期: ", True, DARK, 12),
                (phases["total_months"], True, BLUE, 14),
            ])
        doc.add_page_break()

    def _add_phase_breakdown(self, doc):
        """阶段级工作量明细"""
        wl = self.result.workload_estimate
        breakdown = wl.get("phase_breakdown", {})
        if not breakdown:
            return
        add_heading(doc, "5.1  阶段级工作量明细", level=2)
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        make_header_row(table, ["阶段", "占比", "人月", "主要工作"])
        for key in sorted(breakdown.keys()):
            ph = breakdown[key]
            add_data_row(table, [
                ph["name"],
                f"{ph['pct']}%",
                f"{round(ph['days']/22, 1)}",
                ph["detail"],
            ])
        doc.add_paragraph()
        doc.add_page_break()

    def _add_capacity_planning(self, doc):
        """容量规划与硬件配置建议"""
        cap = self.result.capacity_planning
        if not cap or "error" in cap:
            return
        add_heading(doc, "6  容量规划与硬件配置建议", level=1)
        doc.add_paragraph()
        add_heading(doc, "6.1  推荐集群配置", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        make_header_row(table, ["项目", "推荐方案"])
        for k, v in [
            ("推荐规格", cap.get("spec", "—")),
            ("CPU", cap.get("cpu", "—")),
            ("内存", cap.get("memory", "—")),
            ("磁盘", cap.get("disk", "—")),
            ("数据节点数", str(cap.get("data_nodes", "—"))),
            ("管理节点数", str(cap.get("manager_nodes", "—"))),
            ("备用节点数", str(cap.get("standby_nodes", "—"))),
            ("总节点数", str(cap.get("total_nodes", "—"))),
        ]:
            add_data_row(table, [k, v])
        doc.add_paragraph()

        add_heading(doc, "6.2  存储估算", level=2)
        sd = cap.get("storage_detail", {})
        table2 = doc.add_table(rows=1, cols=2)
        table2.style = 'Table Grid'
        make_header_row(table2, ["项目", "估算值"])
        for k, v in [
            ("当前数据量", f"{sd.get('current_tb', '—')} TB"),
            ("5年后数据量", f"{sd.get('future_5yr_tb', '—')} TB"),
            ("压缩后(5:1)", f"{sd.get('compressed_tb', '—')} TB"),
            ("含系统开销(总需求)", f"{sd.get('total_with_overhead_tb', '—')} TB"),
        ]:
            add_data_row(table2, [k, v])
        doc.add_paragraph()

        if cap.get("concurrency_advice"):
            add_heading(doc, "6.3  并发建议", level=2)
            add_para(doc, cap["concurrency_advice"], size=10)

        if cap.get("notes"):
            add_heading(doc, "6.4  说明", level=2)
            for note in cap["notes"]:
                add_para(doc, f"• {note}", color=GRAY, size=9)
        doc.add_page_break()

    def _add_batch_strategy(self, doc):
        """分批迁移策略建议"""
        bs = self.result.batch_strategy
        if not bs:
            return
        add_heading(doc, "7  分批迁移策略", level=1)
        doc.add_paragraph()
        add_para(doc, f"策略: {bs.get('strategy', '—')}", bold=True, size=11)
        doc.add_paragraph()

        batches = bs.get("batches", [])
        if batches:
            add_heading(doc, "7.1  分批方案", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            make_header_row(table, ["批次", "迁移范围", "窗口"])
            for b in batches:
                add_data_row(table, [
                    f"第{b['batch']}批",
                    b.get("scope", "—"),
                    b.get("window", "—"),
                ])
            doc.add_paragraph()

        precheck = bs.get("precheck", [])
        if precheck:
            add_heading(doc, "7.2  执行前检查清单", level=2)
            for item in precheck:
                add_para(doc, f"☐ {item}", size=9)
        doc.add_page_break()

    def _add_recommendations(self, doc):
        add_heading(doc, "9  建议与推荐", level=1)
        doc.add_paragraph()
        recs = self.result.recommendations
        if not recs:
            add_para(doc, "无特殊建议", color=GRAY)
            return
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        make_header_row(table, ["序号", "建议内容"])
        for i, rec in enumerate(recs, 1):
            add_data_row(table, [str(i), rec])
        doc.add_paragraph()
        add_heading(doc, "9.1  推荐迁移工具链", level=2)
        t2 = doc.add_table(rows=1, cols=3)
        t2.style = 'Table Grid'
        make_header_row(t2, ["领域", "推荐工具", "说明"])
        for domain, tool, note in [
            ("DDL转换", "华为UGO + DSC + 脚本批量转换", "自动语法分析+人工复核"),
            ("数据迁移(全量)", "华为DRS / DataX / GDS", "DRS在线全量; DataX灵活配置; GDS百GB级并行导入"),
            ("数据迁移(增量)", "华为DRS / Debezium + Kafka + Flink", "DRS实时同步; CDC实时入仓"),
            ("ETL调度", "DolphinScheduler / TaskCTL", "统一调度平台，支持DWS原生"),
            ("数据校验(L1-L4)", "MD5校验 + 行数对比 + 业务逻辑验证", "L1结构/L2数据量/L3内容/L4业务"),
            ("安全审计迁移", "华为云KMS + DWS安全策略", "TDE密钥迁移/RLS策略重写/AUDIT配置"),
            ("应用层适配", "JDBC/ODBC驱动替换 + ORM方言配置", "Npgsql/EF Core/MyBatis方言适配"),
        ]:
            add_data_row(t2, [domain, tool, note])
        doc.add_paragraph()

        # L1-L4数据一致性校验框架
        add_heading(doc, "9.2  L1-L4四级数据一致性校验框架", level=2)
        t3 = doc.add_table(rows=1, cols=5)
        t3.style = 'Table Grid'
        make_header_row(t3, ["层级", "名称", "校验方法", "范围", "时机"])
        verifications = [
            ("L1", "结构校验", "DDL自动比对: 字段数量/类型/约束/索引", "全部对象", "全量迁移后"),
            ("L2", "数据量校验", "SELECT COUNT(*)/COUNT(DISTINCT PK)", "全部表", "全量+增量后"),
            ("L3", "数据内容校验", "数值SUM/AVG/MIN/MAX + MD5分块抽样", "核心业务表", "全量+定期增量后"),
            ("L4", "业务逻辑校验", "资产总值核对/份额勾稽/T+1清算验证", "按业务场景定制", "试运行1-2周"),
        ]
        for v in verifications:
            add_data_row(t3, list(v))
        doc.add_paragraph()
        add_para(doc, (
            "说明: 建议在迁移全过程中严格执行L1→L2→L3→L4逐级校验体系，"
            "每级校验通过后方可进入下一阶段。L4业务校验是金融行业验收的核心标准。"
        ), color=GRAY, size=9)
        doc.add_page_break()

    def _add_appendix(self, doc):
        add_heading(doc, "附录", level=1)
        doc.add_paragraph()
        add_heading(doc, "附录A  评分规则说明", level=2)
        add_para(doc, (
            "评估采用加权评分模型，总分0-100分:\n"
            "• 90-100分: 低风险，大部分对象可直接迁移\n"
            "• 70-89分: 中风险，存在少量需改造项\n"
            "• 50-69分: 高风险，存在较多不兼容项\n"
            "• <50分: 极高风险，建议谨慎评估迁移可行性\n\n"
            "评分维度权重会根据迁移路径动态分配，核心维度包括: "
            "DDL兼容性、数据类型兼容性、函数兼容性、UDF/存储过程语言兼容性、"
            "扩展兼容性、安全与权限、字符集与编码、应用层兼容性、事务与并发、"
            "CDC增量同步、性能与容量、ETL/调度/BI工具兼容性。\n\n"
            "为确保评估全面性，系统还提供L1-L4四级数据一致性校验框架建议:"
            "\n• L1结构校验: 表数量/字段/约束/索引一致性\n"
            "• L2数据量校验: 行数对比/主键唯一性\n"
            "• L3数据内容校验: MD5抽样/SUM/AVG/MAX/MIN对比\n"
            "• L4业务逻辑校验: 行业特定业务规则验证"
        ), size=9)
        doc.add_paragraph()
        add_heading(doc, "附录B  声明", level=2)
        add_para(doc, (
            "本报告基于调研模板数据和内置兼容性规则库自动生成，"
            "评估结果仅供参考。"
            "实际迁移中需结合具体业务场景、SQL复杂度、性能要求等因素综合判断。"
            "建议在POC阶段对关键业务SQL和UDF进行真实环境验证。"
        ), size=9, color=GRAY)
